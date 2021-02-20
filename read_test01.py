# -*- coding: utf-8 -*-

###########################################################################
## Python code generated with wxFormBuilder (version 3.9.0 Dec  4 2019)
## http://www.wxformbuilder.org/
##
## PLEASE DO *NOT* EDIT THIS FILE!
###########################################################################

import wx
import wx.xrc
import os
###########################################################################
## Class MyFrame1
###########################################################################

class MyFrame1 ( wx.Frame ):

	def __init__( self, parent ):
		wx.Frame.__init__ ( self, parent, id = wx.ID_ANY, title = wx.EmptyString, pos = wx.DefaultPosition, size = wx.Size( 500,300 ), style = wx.DEFAULT_FRAME_STYLE|wx.TAB_TRAVERSAL )

		self.SetSizeHints( wx.DefaultSize, wx.DefaultSize )

		bSizer1 = wx.BoxSizer( wx.VERTICAL )

		self.m_button3 = wx.Button( self, wx.ID_ANY, u"Select File", wx.DefaultPosition, wx.DefaultSize, 0 )
		bSizer1.Add( self.m_button3, 0, wx.ALL, 5 )

		self.m_textCtrl3 = wx.TextCtrl( self,  wx.ID_ANY, wx.EmptyString, wx.DefaultPosition, wx.DefaultSize, wx.TE_READONLY|wx.TE_MULTILINE)
		bSizer1.Add( self.m_textCtrl3, 1, wx.ALL|wx.EXPAND, 5 )

		self.m_button5 = wx.Button( self, wx.ID_ANY, u"Run", wx.DefaultPosition, wx.DefaultSize, 0 )
		bSizer1.Add( self.m_button5, 0, wx.ALL, 5 )

		self.SetSizer( bSizer1 )
		self.Layout()
		self.m_menubar2 = wx.MenuBar( 0 )
		self.SetMenuBar( self.m_menubar2 )

		self.m_statusBar2 = self.CreateStatusBar( 1, wx.STB_SIZEGRIP, wx.ID_ANY )

		self.Centre( wx.BOTH )

		# Connect Events
		self.m_button3.Bind( wx.EVT_BUTTON, self.m_button3OnButtonClick )
		self.m_button5.Bind( wx.EVT_BUTTON, self.m_button5OnButtonClick )		

	def __del__( self ):
		pass


	# Virtual event handlers, overide them in your derived class
	# def m_button3OnButtonClick( self,event ):
	#	 openFileDialog = wx.FileDialog(frame, "请选择要打开的文件", "", "","word格式 (*.docx)|*.docx",wx.FD_OPEN | wx.FD_MULTIPLE)

	# Virtual event handlers, overide them in your derived class
	def m_button3OnButtonClick(self, event):
		openFileDialog = wx.FileDialog(frame, "请选择要打开的文件", "", "",
									   "*.*",
									   wx.FD_OPEN | wx.FD_MULTIPLE)
		if openFileDialog.ShowModal() == wx.ID_OK:
			filePath = openFileDialog.GetDirectory()
			filename=openFileDialog.GetFilenames()
			self.m_textCtrl3.SetValue(filePath+'\n') 
			self.m_textCtrl3.AppendText('\n'.join(filename)) 

	def m_button5OnButtonClick(self, event):
			file_all=self.m_textCtrl3.GetValue()
			file_list=file_all.split('\n')
			files=[]
			docx_num=0
			xlsx_num=0
			for fl in file_list[1:]:
				files.append(os.path.join(file_list[0],fl))
				if fl[-5:]=='.docx':
					docx_num=docx_num+1
				if fl[-5:]=='.xlsx':
					xlsx_num=xlsx_num+1
			if docx_num==len(files):
				self.m_statusBar2.SetStatusText("Reading Mockup beginning ...")
#################################### Start Read Mockup ##################################								
				#  解决header的格式问题：  ##
				# 修改header，如果为default，则不合并第一列header单元格 #

				#识别header结束行并获取行号list
				def read_headline(table):
					headline_flag=[10]
					for n, row in enumerate(table.rows):   # 读每行
						if n < 10:
							for j,cell in enumerate(row.cells):  # 读一行中的所有单元格
								if table.cell(n,0).text.strip() != '':
									if table.cell(n,1).text.find('xx (') >= 0 or table.cell(n,1).text.strip()=='xx':
										headline_flag.append(n-1)   
								if table.cell(n,0).text.strip() == '' or table.cell(n,0).text.strip() != '':
									cellmiss=0
									for i in range(1,len(row.cells)):
										if table.cell(n,i).text.strip()=='':
											cellmiss=cellmiss+1
									if cellmiss==len(row.cells)-1:
										headline_flag.append(n-1) 
								if table.cell(n,0).text.strip().startswith('xxx'):
									headline_flag.append(n-1) 
					headline_flag_all.append(min(headline_flag))  

				import docx
				from docx import Document
				import re
				#mockup位置
				# test_d0 = 'C:/Users/mingfeng.zhou/Desktop/standard_mockup/BeiGene non-efficacy TFL standard_02Sep2019_final_test1.docx'
				# test_d1 = 'C:/Users/mingfeng.zhou/Desktop/standard_mockup/BeiGene non-efficacy TFL standard_02Sep2019_final_test1.docx'
				# test_d2 = 'C:/Users/mingfeng.zhou/Desktop/standard_mockup/BeiGene non-efficacy TFL standard_02Sep2019_final_test1.docx'
				# files=[test_d0,test_d1,test_d2]
				files=files

				#下载的图片本地保存位置
				# output_pictures='C:/Users/mingfeng.zhou/Desktop/standard_mockup/mockup-pictures/'
				output_pictures=file_list[0]+'/mockup-pictures/'

				paragraphlist=[]
				# 读取正文内容：Paragraphs
				# print("段落数:"+str(len(document.paragraphs)))#段落数
				table_num=[]
				title=[]
				population=[]
				footnote=[]
				table_code=[]
				table_all=[]
				footnote_sub=''
				fig_code=[]
				footnote_sub_list=[]
				table_contents_all=[]
				first_col_hd=[]
				first_col_hd_default=[]
				param_all=[]
				headline_flag_all=[]
				table_hd_all=[]
				for fn,test_d in enumerate(files):
					document=Document(test_d)
					table_value=10000
					datasource_num=100000
					for i,paragraph in enumerate(document.paragraphs):
						text= paragraph.text
						code=paragraph._p.getnext()
						paragraphlist.append(text)
						xmlcode=paragraph._p.xml
						run=paragraph.runs
						# print(i)
						#print(text)
						# print(code.tag)
						# print(code)

						if text.startswith(('Table','List','Figure')) and text.rstrip().endswith(')'): 
							table_num.append('file'+str(fn)+'_'+text)
							table_value=i  
							paragraphlist[table_value]='file'+str(fn)+'_'+text
							# print(paragraphlist[table_value])
							tbl_code=''
							rid='' 

						if 'graphicData' in xmlcode:  #查询图片位置
							rid="".join(re.findall(r'a:blip r:embed="(\w+)"',xmlcode))

						if i > table_value and i < table_value + 5 and code.tag.endswith('tbl'):
							tbl_code=code
							if tbl_code!='':
								for m,tbl in enumerate(document.tables):
									if tbl._tbl==tbl_code: # 获取对应位置的table

										self.m_statusBar2.SetStatusText("Reading Mockup "+paragraphlist[table_value]+" ...")

										read_headline(tbl) #获取header最后一行行号
										row_text=[]
										row_id=[]
										row_hd=[]
										first_col_hd_subset=[]
										first_col_hd_subset_default=[]
										for n, row in enumerate(tbl.rows):   # 读每行
											col_text=[]
											col_hd=[]
											first_hd_0=[]
											for j,cell in enumerate(row.cells):  # 读一行中的所有单元格
												cell_run_text=[]
												for z,cell_runs in enumerate(cell.paragraphs):
													cell_run = cell_runs.runs
													for k,cell_p in enumerate(cell_run):
														p_font=[cell_p.text,[cell_p.bold,cell_p.italic,cell_p.underline,cell_p.font.superscript]]   
														cell_run_text.append(p_font)
													if z<len(cell.paragraphs)-1:
														cell_run_text.append(['\n',[None,None,None,None]])

												if j > 0 :
													col_hd.append([cell.text,cell_run_text])

												if j ==0:
													col_text.append([paragraphlist[table_value],cell.text,cell_run_text])
													first_hd_0.append(cell_run_text)
							
												else:
													col_text.append([cell.text,cell_run_text])  

												if n<=headline_flag_all[-1] and j==0:
													first_col_hd_subset.extend(cell_run_text)
													# if n<headline_flag_all[-1]:
													#	 first_col_hd_subset=first_col_hd_subset.extend([['\n',None,None,None,None]])									   
					
											# if row.cells[0].text.strip()!='':
											row_text.append(col_text)
										
											if n <= headline_flag_all[-1]:
												colmiss=0
												for ms in range(0,len(col_hd)):
													if col_hd[ms][0].strip()=='':
														colmiss=colmiss+1
												if colmiss!=len(col_hd):
													row_hd.append(col_hd)
													first_col_hd_subset_default.append(first_hd_0)

											if len(row_text[n])>0:
												if len(row_text[n][0]) > 0 and n > headline_flag_all[-1]:
													if len(row_text[n][0][2])>0:
														if row_text[n][0][2][0][0].lstrip()==row_text[n][0][2][0][0] and paragraphlist[table_value].startswith('file'+str(fn)+'_Table'):
															param_all.append([paragraphlist[table_value],row_text[n][0][1],row_text[n][0][2]])
										table_contents_all.append(row_text)   
										# if paragraphlist[table_value].find('_Listing')==-1:
										table_hd_all.append([paragraphlist[table_value],row_hd])
										first_col_hd.append([paragraphlist[table_value],first_col_hd_subset])
										first_col_hd_default.append([paragraphlist[table_value],first_col_hd_subset_default])
						run_text=[]
						for p in run:
							p_font=[p.text,[p.bold,p.italic,p.underline,p.font.superscript]]
							run_text.append(p_font)

						if i==table_value+1:
							title.append([paragraphlist[table_value],text,run_text])

						if i==table_value+2:
							population.append([paragraphlist[table_value],text,run_text])

						if text.strip().startswith('Data source:'):
							datasource_num=i
							# print(datasource_num)
							footnote_sub=text
							footnote_sub_list=[[text,run_text]]
							# print('____foot_init_',footnote_sub_list)
							table_code.append(tbl_code)  
							fig_code.append([paragraphlist[table_value],rid])

						if i> datasource_num:
							footnote_sub_list.append([text,run_text])
							# print('append',footnote_sub_list)
							footnote_sub=footnote_sub+text
							if text.strip().startswith('Protocol') or text.strip().startswith('Table ') or text.strip().startswith('Listing ') or text.strip().startswith('Figure '):
								#print(footnote_sub_list)
								footnote_sub=footnote_sub.replace('\n'+paragraphlist[i-1],'').replace(paragraphlist[i-1],'').replace(text,'')
								footnote_list=[paragraphlist[table_value],footnote_sub,footnote_sub_list[0:i-1]]
								footnote.append(footnote_list)
								#print(footnote)
					footnote_list=[paragraphlist[table_value],footnote_sub,footnote_sub_list]
					footnote.append(footnote_list)

					#填补headline_flag_all的figure部分的值，默认为0
					len_headline_flag=len(headline_flag_all)
					for i in range(len(table_num)):
						if i >=len_headline_flag:
							headline_flag_all.append(0)

				# 将document中的图片下载到本地
					for shape in document.inline_shapes:
						contentID = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
						# print(contentID)
						contentType = document.part.related_parts[contentID].content_type
						# print (shape.height.cm,shape.width.cm,shape._inline.graphic.graphicData.pic.nvPicPr.cNvPr.name) 
						if not contentType.startswith('image'):
							continue
						imgName = document.part.related_parts[contentID].partname
						imgData = document.part.related_parts[contentID]._blob
						# print(imgName,imgData)
						with open(output_pictures+'pic'+str(fn)+contentID+'.jpg','wb') as fp:
							fp.write(imgData)
					if len(document.inline_shapes) > 0:
						fp.close()


				import pandas as pd
				import numpy as np 
				import xlsxwriter
				# excel_loc='C:/Users/mingfeng.zhou/Desktop/standard_mockup/result1.xlsx'
				excel_loc=os.path.join(file_list[0],'MockupToExcel.xlsx')

				df_title=pd.DataFrame(title,columns=['table_num','title','title_sep'])
				df_population=pd.DataFrame(population,columns=['table_num','population','pop_sep'])
				df_footnote=pd.DataFrame(footnote,columns=['table_num','footnote','footnote_style'])
				df_first_col_hd=pd.DataFrame(first_col_hd,columns=['table_num','first_col_hd'])
				df_first_col_hd_default=pd.DataFrame(first_col_hd_default,columns=['table_num','first_col_hd_default'])
				df_fig_code=pd.DataFrame(fig_code,columns=['table_num','fig_code'])
				df_table_hd_all=pd.DataFrame(table_hd_all,columns=['table_num','hd_contents'])

				df_tab=pd.concat([df_title[['table_num','title']], df_population[['population']], 
							df_footnote[['footnote']],df_fig_code['fig_code']],
							axis=1)
				df_tab=pd.merge(df_tab,df_first_col_hd[['table_num','first_col_hd']],how='left',on=['table_num'])
				df_tab=pd.merge(df_tab,df_table_hd_all,how='left',on=['table_num'])
				df_tab['header']='default'
				df_tab['headerline_flag']=headline_flag_all
				df_tab=pd.merge(df_tab,df_first_col_hd_default[['table_num','first_col_hd_default']],how='left',on=['table_num'])

# 				variable_all=[]
# 				listing_all=[]
# 				non_param_all=[]
# 				for i,tables in enumerate(table_contents_all):
# 					for j,row in enumerate(tables): 

# 						if len(row)>0 and j > headline_flag_all[i]:
# 							if len(row[0]) > 0:
# 								if row[0][1].strip()!='' and row[0][0].strip().find('_Table')>=0:
# 									variable_all.append(row[0])
# 									# non_param_all.append([row[0][0],row])
# 						if row[0][0].strip().find('_Listing')>=0 and j > headline_flag_all[i]:
# 							listing_all.append([row[0][0],row])
# 						if row[0][0].strip().find('_Table')>=0 and j > headline_flag_all[i]:
# 							non_param_all.append([row[0][0],row])

# 				df_variable=pd.DataFrame(variable_all,columns=['table_num','variable','format_style'])
# 				df_param=pd.DataFrame(param_all,columns=['table_num','param','param_style'])
# 				df_param['variable']=df_param['param']
# 				df_variable_all=pd.merge(df_variable,df_param,how='left',on=['table_num','variable'])
# 				df_variable_all['param']=df_variable_all.param.fillna(method='ffill')
# 				df_format=df_variable_all.groupby(['table_num','param']).apply(lambda group: group.iloc[1:, 0:])
# 				df_format['format']=df_format['variable']
# 				df_format=df_format[['table_num','param','format','format_style']].reset_index(drop=True)
# 				df_listing=pd.DataFrame(listing_all,columns=['list_num','row'])
# 				df_non_param=pd.DataFrame(non_param_all,columns=['table_num','row'])


# 				#excel界面呈现：
# 				df_mockup=df_tab[['table_num','title','population','header']]
# 				for i in range(len(df_tab)):
# 					df_mockup.loc[i,'order']=i

# 				def sub_table_num(x):
# 					loc=max(x.find('_Table'),x.find('_List'),x.find('_Figure'))
# 					renew_x=x[loc+1:]
# 					return(renew_x)
# 				df_mockup['new_table_num']=df_mockup['table_num'].apply(sub_table_num)

# 				workbook=xlsxwriter.Workbook(excel_loc)
# 				# df_mockup.to_excel(writer,sheet_name='mockup',index_label='id')
# 				worksheet=workbook.add_worksheet('mockup')
# 				dict_mockup=df_mockup.to_dict(orient='split')
# 				data=dict_mockup['data']
# 				for i in range(len(data)):
# 					data[i].insert(0,dict_mockup['index'][i])
# 				columns=dict_mockup['columns']
# 				columns.insert(0,'id')
# 				col_dict=[]
# 				for col in columns:
# 					dict_tmp={}
# 					dict_tmp['header']=col
# 					col_dict.append(dict_tmp)
# 				worksheet.add_table(0,0,len(df_mockup),6,
# 							{'data':data,
# 							'columns':col_dict,
# 							'style':'Table Style Medium 2'})
# 				worksheet.set_column('B:D',35)
# 				worksheet.set_column('G:G',35)
# 				# worksheet.set_row(0,None,None,{'locked':True})
# 				# worksheet.set_column('A:B',None,None,{'locked':True})
# 				workbook.close()

# 				import openpyxl
# 				book=openpyxl.load_workbook(excel_loc)
# 				writer = pd.ExcelWriter(excel_loc,engine='openpyxl')
# 				writer.book=book
# 				df_tab.to_excel(writer, sheet_name='table')
# 				df_param[['table_num','param','param_style']].to_excel(writer, sheet_name='param_style')
# 				df_format.to_excel(writer, sheet_name='format_style')
# 				df_footnote.to_excel(writer, sheet_name='footnote_style')
# 				df_listing.to_excel(writer,sheet_name='listing_style')
# 				df_non_param.to_excel(writer,sheet_name='non_param_style')
# 				writer.save()

# 				#隐藏sheet
# 				xls_book = openpyxl.load_workbook(filename=excel_loc)
# 				sheet_names = xls_book.sheetnames
# 				for i,sheetname in enumerate(sheet_names):
# 					if sheetname!='mockup':
# 						xls_sheet = xls_book[sheetname]
# 						xls_sheet.sheet_state = 'hidden'
# 				xls_book.save(excel_loc)

# ################################# End #############################


				self.m_statusBar2.SetStatusText("The Process of Creating Mockup is Sucessfully End")
			else:
				self.m_statusBar2.SetStatusText("Please check the types of files extension")				


app = wx.App(False)
frame = MyFrame1(None)
frame.Show()
app.MainLoop()