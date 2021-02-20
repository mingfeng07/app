## Project Name : Mockup Auotmation Development Tools
## Programmer   : Mingfeng zhou
## Version      : 1.0
## Purpose      : Read contents from standard Mockup and Create new Mockup you wanted.
## Introduction : This program is divided into three steps:
##                1. Development a GUI to be the enter interface;
##                2. Read contents from standard Mockup;
##                3. Create new Mockup you wanted.				






################################################################################
#     The first step: Development a GUI to be the enter interface              #
################################################################################

# -*- coding: utf-8 -*-

###########################################################################
## Python code generated with wxFormBuilder (version 3.9.0 Dec  4 2019)
## http://www.wxformbuilder.org/
##
## PLEASE DO *NOT* EDIT THIS FILE!
###########################################################################

import wx
import wx.xrc
import os
import traceback
###########################################################################
## Class MyFrame1
#Modification:Add protocol ID bosizer
###########################################################################

class MyFrame1 ( wx.Frame ):

	def __init__( self, parent ):
		wx.Frame.__init__ ( self, parent, id = wx.ID_ANY, title = wx.EmptyString, pos = wx.DefaultPosition, size = wx.Size( 500,300 ), style = wx.DEFAULT_FRAME_STYLE|wx.TAB_TRAVERSAL )

		self.SetSizeHints( wx.DefaultSize, wx.DefaultSize )

		bSizer1 = wx.BoxSizer( wx.VERTICAL )

		self.m_button3 = wx.Button( self, wx.ID_ANY, u"Select File", wx.DefaultPosition, wx.DefaultSize, 0 )
		bSizer1.Add( self.m_button3, 0, wx.ALL, 5 )

		self.m_textCtrl3 = wx.TextCtrl( self,  wx.ID_ANY, wx.EmptyString, wx.DefaultPosition, wx.DefaultSize, wx.TE_READONLY|wx.TE_MULTILINE)
		bSizer1.Add( self.m_textCtrl3, 1, wx.ALL|wx.EXPAND, 5 )

		self.m_button5 = wx.Button( self, wx.ID_ANY, u"Run", wx.DefaultPosition, wx.DefaultSize, 0 )
		bSizer1.Add( self.m_button5, 0, wx.ALL, 5 )

		self.SetSizer( bSizer1 )
		self.Layout()
		self.m_menubar2 = wx.MenuBar( 0 )
		self.SetMenuBar( self.m_menubar2 )

		self.m_statusBar2 = self.CreateStatusBar( 1, wx.STB_SIZEGRIP, wx.ID_ANY )

		self.Centre( wx.BOTH )

		# Connect Events
		self.m_button3.Bind( wx.EVT_BUTTON, self.m_button3OnButtonClick )
		self.m_button5.Bind( wx.EVT_BUTTON, self.m_button5OnButtonClick )		

	def __del__( self ):
		pass

	# Virtual event handlers, overide them in your derived class
	def m_button3OnButtonClick(self, event):
		openFileDialog = wx.FileDialog(frame, "please select file", "", "",
									   "*.*",
									   wx.FD_OPEN | wx.FD_MULTIPLE)
		if openFileDialog.ShowModal() == wx.ID_OK:
			filePath = openFileDialog.GetDirectory()
			filename=openFileDialog.GetFilenames()
			self.m_textCtrl3.SetValue(filePath+'\n') 
			self.m_textCtrl3.AppendText('\n'.join(filename)) 

	def m_button5OnButtonClick(self, event):
			file_all=self.m_textCtrl3.GetValue()
			file_list=file_all.split('\n')
			files=[]
			docx_num=0
			xlsx_num=0
			for fl in file_list[1:]:
				files.append(os.path.join(file_list[0],fl))
				if fl[-5:]=='.docx':
					docx_num=docx_num+1
				if fl[-5:]=='.xlsx':
					xlsx_num=xlsx_num+1
			if docx_num==len(files):
				self.m_statusBar2.SetStatusText("Reading Mockup beginning ...")

################################################################################
#                             The first step end                               #
################################################################################


################################################################################
#     The second step: Read contents from standard Mockup                      #
################################################################################

#################################### Start Read Mockup ##################################								
				#  solve header format problem：  ##
				# modify header，if default，do not comine first column header cell#
				# add try except statment to display error #
				# solve duplicate table id problem
				# solve merger cell problems in header

				# identify the end line of header and get a list of the line number 
				def read_headline(table):
					headline_flag=[10]
					for n, row in enumerate(table.rows):   # read line
						if n < 10:
							for j,cell in enumerate(row.cells):  # read cell in line
								if table.cell(n,0).text.strip() != '':
									if table.cell(n,1).text.find('xx (') >= 0 or table.cell(n,1).text.strip()=='xx':
										headline_flag.append(n-1)   
								if table.cell(n,0).text.strip() == '' or table.cell(n,0).text.strip() != '':
									cellmiss=0
									for i in range(1,len(row.cells)):
										if table.cell(n,i).text.strip()=='':
											cellmiss=cellmiss+1
									if cellmiss==len(row.cells)-1:
										headline_flag.append(n-1) 
								if table.cell(n,0).text.strip().startswith('xxx'):
									headline_flag.append(n-1) 
					return(min(headline_flag))

				import docx
				from docx import Document
				import re
				#mockup position
				# test_d0 = 'C:/Users/mingfeng.zhou/Desktop/standard_mockup/BeiGene non-efficacy TFL standard_02Sep2019_final_test1.docx'
				# test_d1 = 'C:/Users/mingfeng.zhou/Desktop/standard_mockup/BeiGene non-efficacy TFL standard_02Sep2019_final_test1.docx'
				# test_d2 = 'C:/Users/mingfeng.zhou/Desktop/standard_mockup/BeiGene non-efficacy TFL standard_02Sep2019_final_test1.docx'
				# files=[test_d0,test_d1,test_d2]
				files=files

				#download picture position
				# output_pictures='C:/Users/mingfeng.zhou/Desktop/standard_mockup/mockup-pictures/'
				output_pictures=file_list[0]+'/mockup-pictures/'

				paragraphlist=[]
				# read the contents in paragraphs
				# print("numbers of paragraphs:"+str(len(document.paragraphs)))
				table_num=[]
				title=[]
				population=[]
				footnote=[]
				table_code=[]
				table_all=[]
				footnote_sub=''
				footnote_flag=0 #define a initial value flag for footnote 
				fig_code=[]
				footnote_sub_list=[]
				table_contents_all=[]
				first_col_hd=[]
				first_col_hd_default=[]
				param_all=[]
				headline_flag_all=[]
				table_hd_all=[]
				try:
					for fn,test_d in enumerate(files):
						document=Document(test_d)
						table_value=10000
						datasource_num=100000
						for i,paragraph in enumerate(document.paragraphs):
							text= paragraph.text
							code=paragraph._p.getnext()
							paragraphlist.append(text)
							xmlcode=paragraph._p.xml
							run=paragraph.runs
							# print(i)
							#print(text)
							# print(code.tag)
							# print(code)

							if text.startswith(('Table','List','Figure')) and text.rstrip().endswith(')'): 
								table_value=i  
								if 'file'+str(fn)+'_'+text not in table_num:
									table_num.append('file'+str(fn)+'_'+text)
									paragraphlist[table_value]='file'+str(fn)+'_'+text
								else:
									table_num.append('file'+str(fn)+'_'+text+str(i)+'(TLF)')
									paragraphlist[table_value]='file'+str(fn)+'_'+text+str(i)+'(TLF)'                                    
								# print(paragraphlist[table_value])
								tbl_code=''
								rid='' 

								if footnote_flag==0 and len(table_num)==2: #when the first table's footnote not exist, assign the footnote's value as null
									footnote=[[table_num[0],'\n',[['\n',[['\n',[None,None,None,None]]]]]]]
								elif footnote_flag==0 and len(table_num)>2: #when the second or more table's footnote not exist, assign the footnote's value as null
									footnote.append([table_num[-2],'\n',[['\n',[['\n',[None,None,None,None]]]]]])

								footnote_flag=0  #define a initial value flag for footnote 

							if 'graphicData' in xmlcode:  #identify the location of the graphic 
								rid=str(fn)+"".join(re.findall(r'a:blip r:embed="(\w+)"',xmlcode))
								headline_flag_all.append([paragraphlist[table_value],0])

							if i > table_value and i < table_value + 5 and code.tag.endswith('tbl'):
								tbl_code=code
								if tbl_code!='':
									for m,tbl in enumerate(document.tables):
										if tbl._tbl==tbl_code: # get the location of the table


											self.m_statusBar2.SetStatusText("Reading Mockup "+paragraphlist[table_value]+" ...")

											headline_flag=read_headline(tbl)
											headline_flag_all.append([paragraphlist[table_value],headline_flag]) #the end line number of header
											row_text=[]
											row_id=[]
											row_hd=[]
											first_col_hd_subset=[]
											first_col_hd_subset_default=[]
											for n, row in enumerate(tbl.rows):   # read line
												col_text=[]
												col_hd=[]
												first_hd_0=[]
												for j,cell in enumerate(row.cells):  # read cell in line
													cell_run_text=[]
													for z,cell_runs in enumerate(cell.paragraphs):
														cell_run = cell_runs.runs
														for k,cell_p in enumerate(cell_run):
															p_font=[cell_p.text,[cell_p.bold,cell_p.italic,cell_p.underline,cell_p.font.superscript]]   
															cell_run_text.append(p_font)
														if z<len(cell.paragraphs)-1:
															cell_run_text.append(['\n',[None,None,None,None]])

													if j>0 and 0<n<=headline_flag: #when the second or more column current cell value equal last cell value, assgin current cell value as null
														if tbl.cell(n,j).text==tbl.cell(n-1,j).text and tbl.cell(n,j).text.strip()!="":
															cell.text='\n'
															cell_run_text=[ ['\n',[None,None,None,None]] ]

													if j > 0 :
														col_hd.append([cell.text,cell_run_text])

													if j==0 and n>0: ##when the first column current cell value equal last cell value, assgin current cell value as null
														if tbl.cell(n,0).text==tbl.cell(n-1,0).text and tbl.cell(n,0).text.strip()!="":
															cell.text='\n'
															cell_run_text=[ ['\n',[None,None,None,None]] ]

													if j ==0:
														col_text.append([paragraphlist[table_value],cell.text,cell_run_text])
														first_hd_0.append(cell_run_text)
								
													else:
														col_text.append([cell.text,cell_run_text])  

													if n<=headline_flag and j==0:
														first_col_hd_subset.extend(cell_run_text)
														# if n<headline_flag:
														#	 first_col_hd_subset=first_col_hd_subset.extend([['\n',None,None,None,None]])									   
						
												# if row.cells[0].text.strip()!='':
												row_text.append(col_text)
											
												if n <= headline_flag:
													colmiss=0
													for ms in range(0,len(col_hd)):
														if col_hd[ms][0].strip()=='':
															colmiss=colmiss+1
													if colmiss!=len(col_hd):
														row_hd.append(col_hd)
														first_col_hd_subset_default.append(first_hd_0)

												if len(row_text[n])>0:
													if len(row_text[n][0]) > 0 and n > headline_flag:
														if len(row_text[n][0][2])>0:
															if row_text[n][0][2][0][0].lstrip()==row_text[n][0][2][0][0] and paragraphlist[table_value].startswith('file'+str(fn)+'_Table'):
																param_all.append([paragraphlist[table_value],row_text[n][0][1],row_text[n][0][2]])
											table_contents_all.append(row_text)   
											# if paragraphlist[table_value].find('_Listing')==-1:
											table_hd_all.append([paragraphlist[table_value],row_hd])
											first_col_hd.append([paragraphlist[table_value],first_col_hd_subset])
											first_col_hd_default.append([paragraphlist[table_value],first_col_hd_subset_default])
							run_text=[]
							for p in run:
								p_font=[p.text,[p.bold,p.italic,p.underline,p.font.superscript]]
								run_text.append(p_font)

							if i==table_value+1:
								title.append([paragraphlist[table_value],text,run_text])

							if i==table_value+2:
								population.append([paragraphlist[table_value],text,run_text])

							if text.strip().startswith('Data source:') or text.strip().startswith('Refer to'):
								datasource_num=i
								protocol_flag=0
								footnote_flag=1 #if text start with Data source, then footnote exist
								# print(datasource_num)
								footnote_sub=text
								footnote_sub_list=[[text,run_text]]
								# print('____foot_init_',footnote_sub_list)
								table_code.append(tbl_code)  
								fig_code.append([paragraphlist[table_value],rid])

							if i> datasource_num:
								footnote_sub_list.append([text,run_text])
								# print('append',footnote_sub_list)
								footnote_sub=footnote_sub+text

								if footnote_flag==1:
									if text.strip().startswith('Protocol'):
										protocol_flag=1
										#print(footnote_sub_list)
										footnote_sub=footnote_sub.replace('\n'+paragraphlist[i-1],'').replace(paragraphlist[i-1],'').replace(text,'')
										if footnote_sub_list[-2][0].strip().startswith('BeiGene'):
											footnote_list=[paragraphlist[table_value],footnote_sub,footnote_sub_list[0:len(footnote_sub_list)-2]] 
										else:
											footnote_list=[paragraphlist[table_value],footnote_sub,footnote_sub_list[0:len(footnote_sub_list)-1]]
										footnote.append(footnote_list)
										#print(footnote)

									elif text.startswith(('Table','List','Figure')) and text.rstrip().endswith(')'):
										if protocol_flag==0:		
											#print(footnote_sub_list)
											footnote_sub=footnote_sub.replace('\n'+paragraphlist[i-1],'').replace(paragraphlist[i-1],'').replace(text,'')
											footnote_list=[table_num[-2],footnote_sub,footnote_sub_list[0:len(footnote_sub_list)-1]]
											footnote.append(footnote_list)
											#print(footnote)
						if footnote_flag==1 or fn==len(files)-1: #when read the final table/listing/figure, fill the final footnote
							footnote_list=[paragraphlist[table_value],footnote_sub,footnote_sub_list]
							footnote.append(footnote_list)

					# download the graphics to the local file
						for shape in document.inline_shapes:
							contentID = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
							# print(contentID)
							contentType = document.part.related_parts[contentID].content_type
							# print (shape.height.cm,shape.width.cm,shape._inline.graphic.graphicData.pic.nvPicPr.cNvPr.name) 
							if not contentType.startswith('image'):
								continue
							imgName = document.part.related_parts[contentID].partname
							imgData = document.part.related_parts[contentID]._blob
							# print(imgName,imgData)
							with open(output_pictures+'pic'+str(fn)+contentID+'.jpg','wb') as fp:
								fp.write(imgData)
						if len(document.inline_shapes) > 0:
							fp.close()

				except Exception as e:
					print(traceback.format_exc())
					self.m_statusBar2.SetStatusText("Finding Error While"+"Reading Mockup "+paragraphlist[table_value]+" ...")

				else:
					import pandas as pd
					import numpy as np 
					import xlsxwriter
					# excel_loc='C:/Users/mingfeng.zhou/Desktop/standard_mockup/result1.xlsx'

					try:
						excel_loc=os.path.join(file_list[0],'MockupToExcel.xlsx')

						df_title=pd.DataFrame(title,columns=['table_num','title','title_sep'])
						df_population=pd.DataFrame(population,columns=['table_num','population','pop_sep'])
						df_footnote=pd.DataFrame(footnote,columns=['table_num','footnote','footnote_style'])
						df_first_col_hd=pd.DataFrame(first_col_hd,columns=['table_num','first_col_hd'])
						df_first_col_hd_default=pd.DataFrame(first_col_hd_default,columns=['table_num','first_col_hd_default'])
						df_fig_code=pd.DataFrame(fig_code,columns=['table_num','fig_code'])
						df_table_hd_all=pd.DataFrame(table_hd_all,columns=['table_num','hd_contents'])
						df_headline_flag_all=pd.DataFrame(headline_flag_all,columns=['table_num','headerline_flag'])

						df_tab=pd.concat([df_title[['table_num','title']], df_population[['population']], 
									df_footnote[['footnote']],df_fig_code['fig_code']],
									axis=1)
						df_tab=pd.merge(df_tab,df_first_col_hd[['table_num','first_col_hd']],how='left',on=['table_num'])
						df_tab=pd.merge(df_tab,df_table_hd_all,how='left',on=['table_num'])
						df_tab['header']='default'
						df_tab=pd.merge(df_tab,df_headline_flag_all,how='left',on=['table_num'])
						df_tab=pd.merge(df_tab,df_first_col_hd_default[['table_num','first_col_hd_default']],how='left',on=['table_num'])

						variable_all=[]
						listing_all=[]
						non_param_all=[]
						for i,tables in enumerate(table_contents_all):
							header_value=int(df_tab[df_tab['table_num']==tables[0][0][0]]['headerline_flag'])
							for j,row in enumerate(tables): 

								if len(row)>0 and j > header_value:
									if len(row[0]) > 0:
										if row[0][1].strip()!='' and row[0][0].strip().find('_Table')>=0:
											variable_all.append(row[0])
											# non_param_all.append([row[0][0],row])
								if row[0][0].strip().find('_Listing')>=0 and j > header_value:
									listing_all.append([row[0][0],row])
								if row[0][0].strip().find('_Table')>=0 and j > header_value:
									non_param_all.append([row[0][0],row])

						df_variable=pd.DataFrame(variable_all,columns=['table_num','variable','format_style'])
						df_param=pd.DataFrame(param_all,columns=['table_num','param','param_style'])
						df_param['variable']=df_param['param']
						df_variable_all=pd.merge(df_variable,df_param,how='left',on=['table_num','variable'])
						df_variable_all['param']=df_variable_all.param.fillna(method='ffill')
						df_format=df_variable_all.groupby(['table_num','param']).apply(lambda group: group.iloc[1:, 0:])
						df_format['format']=df_format['variable']
						df_format=df_format[['table_num','param','format','format_style']].reset_index(drop=True)
						df_listing=pd.DataFrame(listing_all,columns=['list_num','row'])
						df_non_param=pd.DataFrame(non_param_all,columns=['table_num','row'])


						#display the excel interface：
						df_mockup=df_tab[['table_num','title','population','header']]
						for i in range(len(df_tab)):
							df_mockup.loc[i,'order']=i

						def sub_table_num(x):
							loc=max(x.find('_Table'),x.find('_List'),x.find('_Figure'))
							renew_x=x[loc+1:]
							return(renew_x)
						df_mockup['new_table_num']=df_mockup['table_num'].apply(sub_table_num)

						workbook=xlsxwriter.Workbook(excel_loc)
						# df_mockup.to_excel(writer,sheet_name='mockup',index_label='id')
						worksheet=workbook.add_worksheet('mockup')
						dict_mockup=df_mockup.to_dict(orient='split')
						data=dict_mockup['data']
						for i in range(len(data)):
							data[i].insert(0,dict_mockup['index'][i])
						columns=dict_mockup['columns']
						columns.insert(0,'id')
						col_dict=[]
						for col in columns:
							dict_tmp={}
							dict_tmp['header']=col
							col_dict.append(dict_tmp)
						worksheet.add_table(0,0,len(df_mockup),6,
									{'data':data,
									'columns':col_dict,
									'style':'Table Style Medium 2'})
						worksheet.set_column('B:D',35)
						worksheet.set_column('G:G',35)
						# worksheet.set_row(0,None,None,{'locked':True})
						# worksheet.set_column('A:B',None,None,{'locked':True})
						workbook.close()

						import openpyxl
						book=openpyxl.load_workbook(excel_loc)
						writer = pd.ExcelWriter(excel_loc,engine='openpyxl')
						writer.book=book
						df_tab.to_excel(writer, sheet_name='table')
						df_param[['table_num','param','param_style']].to_excel(writer, sheet_name='param_style')
						df_format.to_excel(writer, sheet_name='format_style')
						df_footnote.to_excel(writer, sheet_name='footnote_style')
						df_listing.to_excel(writer,sheet_name='listing_style')
						df_non_param.to_excel(writer,sheet_name='non_param_style')
						writer.save()

						#hide the sheet you wanted
						xls_book = openpyxl.load_workbook(filename=excel_loc)
						sheet_names = xls_book.sheetnames
						for i,sheetname in enumerate(sheet_names):
							if sheetname!='mockup':
								xls_sheet = xls_book[sheetname]
								xls_sheet.sheet_state = 'hidden'
						xls_book.save(excel_loc)

					except Exception as e:
						print(traceback.format_exc())
						self.m_statusBar2.SetStatusText("Finding Error Before Output Data To Excel...")

################################# End #############################
					else:
						self.m_statusBar2.SetStatusText("The Process of Reading Mockup is Sucessfully End")		

############################## Start Create Mockup ######################################

			elif xlsx_num==1 and len(files)<=2:
				for fl in files:
					if '.xlsx' in fl:
						excel_str=fl
					elif '.docx' in fl:
						word_str=fl

				self.m_statusBar2.SetStatusText("Creating Mockup From Excel beginning ...")

				#####   solve the format of header while reading   #####
				# modify the header, if default then do not merge the first column cell#
				# if exist gap column in header, drop column in relevant body part and keep gap column in small width#

				from docx import Document
				import pandas as pd
				import numpy as np 
				from docx.oxml import OxmlElement
				from docx.oxml.ns import qn

				from docx.enum.text import WD_ALIGN_PARAGRAPH
				from docx.enum.table import WD_ALIGN_VERTICAL
				from docx.shared import Pt
				from docx.shared import RGBColor
				from docx.shared import Inches
				from docx.enum.section import WD_ORIENT
				from docx.enum.section import WD_SECTION
				from docx.enum.table import WD_ROW_HEIGHT_RULE
				from docx.enum.style import WD_STYLE_TYPE
				from docx.enum.text import WD_TAB_ALIGNMENT
				import ast
				import math
				# get header from new docx 
				def read_header_from_docx(path):
					document=Document(path)
					# doc_tb=Document()
					header_list=[]
					for paragraph in document.paragraphs:
						text= paragraph.text
						if 'Study' in text:
							header_list.append(text)
					header_contents_all=dict()
					single_cell_underline_all=dict()
					gap_col_all=dict()
					for m,tbl in enumerate(document.tables):
						row_text=[]
						single_cell_underline_list=[]
						gap_col_list=[]
						for n, row in enumerate(tbl.rows):  
							col_text=[]
							for j,cell in enumerate(row.cells):  

								# read gap column
								if j < len(row.cells)-1 and j > 0 and n < len(tbl.rows)-1:
									if tbl.rows[n].cells[j].text=='' and tbl.rows[n+1].cells[j].text=='':
										if j==1 and tbl.rows[n].cells[j+1].text==tbl.rows[n].cells[j+2].text:
											if j+1 not in gap_col_list:
												gap_col_list.append(j+1)
										elif j > 1:
											if (tbl.rows[n].cells[j+1].text==tbl.rows[n].cells[j+2].text) or (tbl.rows[n].cells[j-2].text==tbl.rows[n].cells[j-1].text):
												if j+1 not in gap_col_list:
													gap_col_list.append(j+1)
								
								cell_run_text=[]  
								for z,cell_runs in enumerate(cell.paragraphs):
									cell_run = cell_runs.runs
									for k,cell_p in enumerate(cell_run):
										p_font=[cell_p.text,[cell_p.bold,cell_p.italic,cell_p.underline,cell_p.font.superscript]]   
										cell_run_text.append(p_font)
									if z==0 and len(cell.paragraphs)>1:
										cell_run_text.append(['\n',[None,None,None,None]])
								col_text.append([cell.text,cell_run_text])							 
							row_text.append(col_text)
						header_contents_all[header_list[m]]=row_text

						#read gap column
						gap_col_all[header_list[m]]=gap_col_list
						print(gap_col_all)
						#read single cell underline from new docx
						tbl0=tbl._tbl 
						for i,cell in enumerate(tbl0.iter_tcs()):  
							tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
							tcBorders = OxmlElement('w:tcBorders')
							bottom = OxmlElement('w:bottom')
							underline=tcPr.first_child_found_in("w:tcBorders")
							if underline is not None:
								single_cell_underline_list.append(i)	
						single_cell_underline_all[header_list[m]]=single_cell_underline_list	 
					return (header_contents_all,single_cell_underline_all,gap_col_all)


				

				#define title, contain：table number，title，population
				def title_set(tableset):
					run1=document.add_heading(tableset.iloc[i,9], level=1)
					run1.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
					run1.paragraph_format.space_before=Pt(5)	 #set the value for space before paragraph 
					run1.paragraph_format.space_after=Pt(5)	     #set the value for space after paragraph
					run1.paragraph_format.line_spacing=Pt(8)	 #set the value for space between line 
					run2=document.add_heading(tableset.iloc[i,5], level=2)  
					run2.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER 
					run2.paragraph_format.space_before=Pt(5)	#set the value for space before paragraph 
					run2.paragraph_format.space_after=Pt(5)	    #set the value for space after paragraph
					run2.paragraph_format.line_spacing=Pt(8)	#set the value for space between line 
					run3=document.add_paragraph(tableset.iloc[i,6])
					run3.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER 
					run3.paragraph_format.space_before=Pt(5)	 #set the value for space before paragraph 
					run3.paragraph_format.space_after=Pt(5)	     #set the value for space after paragraph
					run3.paragraph_format.line_spacing=Pt(8)	 #set the value for space between line 

				#set the style of cell, return the list of merger cell and endlien of header
				def header_cellstyle_set(header_list,tab,default):
					row_header_cnt=len(header_list)
					col_header_cnt=len(header_list[0]) 
					cell_merge_of_list=[]
					header_lastline_cell_list=[]
					header_firstline_cell_list=[]
					dict_value=0   
					for n in range(row_header_cnt):
						for k in range(col_header_cnt+1):
							if k==0 and n>0:
								if default==0:
									row_up,row_down=tab.cell(n-1,k),tab.cell(n,k)
									row_up.merge(row_down)
								dict_value=dict_value+1
							elif k>1:
								if header_list[n][k-2][0]==header_list[n][k-1][0] and header_list[n][k-1][0]!="" and n==0: 
									col_lt,col_rt=tab.cell(n,k-1),tab.cell(n,k)
									col_lt.merge(col_rt)	
									cell_merge_of_list.append(dict_value)
								elif header_list[n][k-2][0]==header_list[n][k-1][0] and header_list[n][k-1][0]!="" and n>0:
									if header_list[n-1][k-2][0]==header_list[n-1][k-1][0] and header_list[n-1][k-1][0].strip().upper()!='(N = XX)':			  
										col_lt,col_rt=tab.cell(n,k-1),tab.cell(n,k)
										col_lt.merge(col_rt)	
										cell_merge_of_list.append(dict_value)		

									elif header_list[n][k-1][0].strip().upper().find('WORST GRADE')>=0:			  
										col_lt,col_rt=tab.cell(n,k-1),tab.cell(n,k)
										col_lt.merge(col_rt)	
										cell_merge_of_list.append(dict_value)

									elif n>1 and header_list[n-1][k-2][0]==header_list[n-1][k-1][0] and header_list[n-1][k-1][0].strip().upper()=='(N = XX)':
										if header_list[n-2][k-2][0]==header_list[n-2][k-1][0]:  
											col_lt,col_rt=tab.cell(n,k-1),tab.cell(n,k)
											col_lt.merge(col_rt)	
											cell_merge_of_list.append(dict_value)							
										else:
											dict_value=dict_value+1
									else:
										dict_value=dict_value+1
								else:
									dict_value=dict_value+1
							elif k==1:
								dict_value=dict_value+1
							if n==row_header_cnt-1:
								header_lastline_cell_list.append(dict_value)
							if n==0:
								header_firstline_cell_list.append(dict_value)
					return(cell_merge_of_list,header_firstline_cell_list,header_lastline_cell_list)

				#set the border line for cell 
				def border_set(tab,header_firstline,header_mid,header_lastline,body_lastline):
					tbl=tab._tbl 
					for i,cell in enumerate(tbl.iter_tcs()):
						if i in header_firstline:
							tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
							tcBorders = OxmlElement('w:tcBorders')
							top = OxmlElement('w:top')
							top.set(qn('w:val'), 'thick')
							tcBorders.append(top)
							tcPr.append(tcBorders)				
						if i in header_firstline and i in header_mid:
							tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
							tcBorders = OxmlElement('w:tcBorders')
							top = OxmlElement('w:top')
							top.set(qn('w:val'), 'thick')
							bottom = OxmlElement('w:bottom')
							bottom.set(qn('w:val'), 'thick')			
							tcBorders.append(top)
							tcBorders.append(bottom)
							tcPr.append(tcBorders)
						elif i in header_mid or i in header_lastline or i in body_lastline:
							tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
							tcBorders = OxmlElement('w:tcBorders')
							bottom = OxmlElement('w:bottom')
							bottom.set(qn('w:val'), 'thick')
							tcBorders.append(bottom)
							tcPr.append(tcBorders)			  

				#set value and style for cell:
				def cell_set(p,cell_list):
					for m in range(len(cell_list)):
						run=p.add_run(cell_list[m][0])
						run.bold=cell_list[m][1][0]
						run.italic=cell_list[m][1][1]
						run.underline=cell_list[m][1][2]
						run.font.superscript=cell_list[m][1][3]	 


				#set value for header cell except for first column header:
				def set_value_header(tab,header_list):
					row=len(header_list)
					col=len(header_list[0])
					for i in range(row):
						for j in range(col):		  
							cell_list=header_list[i][j][1]
							if j>0:
								if header_list[i][j][0]!=header_list[i][j-1][0]:	
								# if header_list[i][j]!=header_list[i][j-1]:	  
									p=tab.cell(i,j+1).paragraphs[0]
									cell_set(p,cell_list)
									# p.add_run(cell_list)
								elif header_list[i][j][0]==header_list[i][j-1][0] and (header_list[i][j][0].strip().upper()=='(N = XX)' or header_list[i][j][0].strip().upper()=='N/N (%)'):
									if i>0:
										if header_list[i-1][j][0]!=header_list[i-1][j-1][0]:
											p=tab.cell(i,j+1).paragraphs[0]
											cell_set(p,cell_list)
											# p.add_run(cell_list)
								elif header_list[i][j][0]==header_list[i][j-1][0] and header_list[i][j][0].strip()=='n (%)':
									if i>0:
										if header_list[i-1][j][0]!=header_list[i-1][j-1][0]:
											p=tab.cell(i,j+1).paragraphs[0]
											cell_set(p,cell_list)
											# p.add_run(cell_list)							 
									if i>1:
										if header_list[i-1][j][0]==header_list[i-1][j-1][0] and header_list[i-1][j][0].strip().upper()=='(N = XX)':
											if header_list[i-2][j][0]!=header_list[i-2][j-1][0]:
												p=tab.cell(i,j+1).paragraphs[0]
												cell_set(p,cell_list)
												# p.add_run(cell_list)					
							else:
								p=tab.cell(i,j+1).paragraphs[0]
								cell_set(p,cell_list)
								# p.add_run(cell_list)
							p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
							p.paragraph_format.space_before=Pt(2)	 
							p.paragraph_format.space_after=Pt(2)	
							p.paragraph_format.line_spacing=Pt(8)	 

				#define general variable
				# header_loc = 'C:/Users/mingfeng.zhou/Desktop/standard_mockup/header_for_tables.docx'
				header_loc=word_str
				# mp_excel_loc='C:/Users/mingfeng.zhou/Desktop/standard_mockup/result1.xlsx'
				mp_excel_loc=excel_str
				# pic_loc= 'C:/Users/mingfeng.zhou/Desktop/standard_mockup/BeiGene_pic.png'
				# fig_read_loc= 'C:/Users/mingfeng.zhou/Desktop/standard_mockup/mockup-pictures/'
				fig_read_loc=file_list[0]+'/mockup-pictures/'
				# output_loc='C:/Users/mingfeng.zhou/Desktop/standard_mockup/demo_output.docx'
				output_loc=os.path.join(file_list[0],'Mockup-Output.docx')
				# Protocol="BGB-3111-210"
				# version="1.6"
				# draftdate="12 Aug 2019"

				doc_header=Document(header_loc)
				document = Document()

				sections=document.sections
				section = document.sections[-1]
				#set orientation：
				section.orientation = WD_ORIENT.LANDSCAPE
				section.page_height,section.page_width=Inches(8.5),Inches(11)
				#section.left_margin, section.right_margin，section.top_margin, section.bottom_margin，section.header_distance, section.footer_distance，section.page_width, section.page_height

				#set font type and size：
				document.styles['Normal'].font.name = u'Courier New'   
				document.styles['Normal'].font.size=Pt(8)

				#set the title style：
				document.styles['Heading 1'].font.size = Pt(8)
				document.styles['Heading 1'].font.name = u'Courier New'
				document.styles['Heading 1'].font.color.rgb = None
				document.styles['Heading 1'].font.italic = False
				document.styles['Heading 1'].font.bold = False
				document.styles['Heading 2'].font.size = Pt(8)
				document.styles['Heading 2'].font.name = u'Courier New'
				document.styles['Heading 2'].font.color.rgb = None
				document.styles['Heading 2'].font.italic = False
				document.styles['Heading 2'].font.bold = False

				table_ori=pd.read_excel(mp_excel_loc,sheet_name='table')
				table_mp=pd.read_excel(mp_excel_loc,sheet_name='mockup')
				table_mp=table_mp[table_mp.order>=0].sort_values('order')
				table=pd.merge(table_ori[['table_num','footnote','fig_code','first_col_hd','hd_contents']],
								table_mp[['table_num','title','population','header','order','new_table_num']],
								how='right',on=['table_num'],sort=False)

				#merge original docx header and new docx header 
				header_list_all,single_cell_underline_all,gap_col_all=read_header_from_docx(header_loc)
				for i in range(len(table)):
					if table.iloc[i,7]=='default' and table.iloc[i,0].find('_Fig')==-1:
						table.iloc[i,7]=table.iloc[i,0]
						if table.iloc[i,4]==table.iloc[i,4]:
							header_list_all[table.iloc[i,0]]=ast.literal_eval(table.iloc[i,4])

				df_param=pd.read_excel(mp_excel_loc,sheet_name='param_style')
				df_non_param=pd.read_excel(mp_excel_loc,sheet_name='non_param_style')
				df_format=pd.read_excel(mp_excel_loc,sheet_name='format_style')
				df_footnote=pd.read_excel(mp_excel_loc,sheet_name='footnote_style')
				table_T=table[table.table_num.str.find('_Table')>-1]
				table_F=table[table.table_num.str.find('_Fig')>-1]
				table_L=table[table.table_num.str.find('_List')>-1]
				df_listing=pd.read_excel(mp_excel_loc,sheet_name='listing_style')



				try:
					#set the value for table
					for i in range(len(table_T)):

						self.m_statusBar2.SetStatusText("Creating Table From Excel: "+table_T.iloc[i,9])
						
						#check whether table exists
						if table_T.iloc[i,4]!=table_T.iloc[i,4]:
							tbl_flag=0
						else:
							tbl_flag=1

						table_num=table_T.iloc[i,0]
						header=table_T.iloc[i,7]

						if tbl_flag==1:
							header_list=header_list_all[header]

						first_col_hd=table_T.iloc[i,3]
						first_col_hd_default=list(table_ori[table_ori['table_num']==table_num]['first_col_hd_default'])
						footnote=table_T.iloc[i,1]
						df_param_subset=df_param[df_param.table_num==table_num]
						df_format_subset=df_format[df_format.table_num==table_num]
						df_footnote_subset=df_footnote[df_footnote.table_num==table_num]
						df_non_param_subset=df_non_param[df_non_param.table_num==table_num]

						#read title, contain：table number，title，population
						title_set(table_T)

						if tbl_flag==1:
							row_header_cnt=len(header_list)
							col_header_cnt=len(header_list[0])
							row_body_cnt=len(df_non_param_subset)
							tab =document.add_table(rows=row_body_cnt+row_header_cnt,cols=1+col_header_cnt)  
							
							#set the merger of cell in header 
							if header.startswith('file'):
								default=1
								single_cell_underline_list=[]
								gap_col_list=[]
							else:
								default=0
								single_cell_underline_list=single_cell_underline_all[header] 
								gap_col_list=gap_col_all[header]

							cell_merge_list,header_firstline_cell_list,header_lastline_cell_list=header_cellstyle_set(header_list,tab,default)   

							print('cellmerge=======',cell_merge_list)
							print('cellsingle=======',single_cell_underline_list)	
							print('cell_header_lastline====',header_lastline_cell_list)
							start=header_lastline_cell_list[-1]+(row_body_cnt-1)*(col_header_cnt+1)+1
							end=header_lastline_cell_list[-1]+row_body_cnt*(col_header_cnt+1)+1
							body_lastline_cell_list=[s for s in range(start,end)]

							#set the border
							border_set(tab,header_firstline_cell_list,cell_merge_list,header_lastline_cell_list,body_lastline_cell_list)

							#set value for header, and center the cell
							set_value_header(tab,header_list)
								#set value the first column for heaer
							if default==0:
								p=tab.cell(row_header_cnt-1,0).paragraphs[0]
								p.paragraph_format.space_before=Pt(2)	 
								p.paragraph_format.space_after=Pt(2)	   
								p.paragraph_format.line_spacing=Pt(8)	   
								first_col_hd_tolist=ast.literal_eval(first_col_hd)
								cell_set(p,first_col_hd_tolist)
							else:
								first_col_hd_tolist=ast.literal_eval(first_col_hd_default[0])
								len_first_col_hd=len(first_col_hd_tolist)
								for i,first_col in enumerate(first_col_hd_tolist):
									p=tab.cell(row_header_cnt-len_first_col_hd+i,0).paragraphs[0]
									p.paragraph_format.space_before=Pt(2)	 
									p.paragraph_format.space_after=Pt(2)	   
									p.paragraph_format.line_spacing=Pt(8)	   
									if first_col[0]!=[]:
										cell_set(p,first_col[0])
							#set value for body
							row_num=row_header_cnt	
							for si,row in enumerate(list(df_non_param_subset['row'])):
								row_list=ast.literal_eval(row)
								len_row=len(row_list)
								for j,col in enumerate(body_lastline_cell_list):
									p=tab.cell(row_num+si,j).paragraphs[0]
									if j==0:
										p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT
									else:
										p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER				
									p.paragraph_format.space_before=Pt(2)	
									p.paragraph_format.space_after=Pt(2)	   
									p.paragraph_format.line_spacing=Pt(8)	  
									if j==0:
										cell_liststyle=row_list[j][1:][1]
									elif j<len_row:
										cell_liststyle=row_list[j][1]
									else:
										cell_liststyle=row_list[len_row-1][1]					
									cell_set(p,cell_liststyle)  
									# gap_column set null
									if j in gap_col_list:
										tab.cell(row_num+si,j).text=''

							#set width for columns
							tab.autofit = False
							col = tab.columns[0] 

							#gap width set as 0.2inches
							gap_width=0.2
							gap_cnt=len(gap_col_list)

							col_cnt=col_header_cnt+1
							if col_cnt <=4:
								width_col0=4
							else:
								width_col0=4-(col_cnt-4)*0.5 
							width_col_oth=(8.5-width_col0-gap_width*gap_cnt)/(col_cnt-1-gap_cnt)
							for ilen in range(col_cnt):
								col = tab.columns[ilen] 
								if ilen==0:
									for cell in col.cells:
										cell.width = Inches(width_col0)
								elif ilen in gap_col_list:
									for cell in col.cells:
										cell.width = Inches(gap_width)
								else:
									for cell in col.cells:
										cell.width = Inches(width_col_oth)		   
		  
						#set value for footnote
						footnote_style=df_footnote_subset.iloc[0,3]
						footnote_style=ast.literal_eval(footnote_style)
						p=document.add_paragraph()
						bullet_flag=0
						if len(footnote_style)>0:
							for m in range(len(footnote_style)):
								if len(footnote_style[m])>0:
									if len(footnote_style[m][1])>0 and footnote_style[m][0].startswith('BeiGene')==0 and footnote_style[m][0].startswith('Protocol BGB')==0: #check whether footnote is null
										if bullet_flag==1 and footnote_style[m][0].strip()!='':
											footnote_style[m][1][0][0]='  ●  '+footnote_style[m][1][0][0]
										cell_set(p,footnote_style[m][1]) 
										if footnote_style[m][0].strip().upper().startswith('PROGRAMMING NOTE'):
											bullet_flag=1
								p.add_run('\n')  
						if tbl_flag==1: 
							document.add_page_break() 

				except Exception as e:
					print(traceback.format_exc())
					self.m_statusBar2.SetStatusText("Finding Error While Creating Table From Excel: "+table_T.iloc[i,9])  
				else:

					try:
						#set value for listing
						for i in range(len(table_L)):

							self.m_statusBar2.SetStatusText("Creating Listing From Excel: "+table_L.iloc[i,9])  

							#check whether table exists
							if table_L.iloc[i,4]!=table_L.iloc[i,4]:
								tbl_flag=0
							else:
								tbl_flag=1

							table_num=table_L.iloc[i,0]
							footnote=table_L.iloc[i,1]
							header=table_L.iloc[i,7]

							if tbl_flag==1:
								header_list=header_list_all[header]

							first_col_hd=table_L.iloc[i,3]
							first_col_hd_default=list(table_ori[table_ori['table_num']==table_num]['first_col_hd_default'])
							headerline_flag=table_ori[table_ori.table_num==table_num]['headerline_flag']
							df_listing_subset=df_listing[df_listing.list_num==table_num]
							df_footnote_subset=df_footnote[df_footnote.table_num==table_num]

							#read title, contain：table number，title，population
							title_set(table_L)

							if tbl_flag==1:
								row_header_cnt=len(header_list)
								col_header_cnt=len(header_list[0])
								row_body_cnt=len(df_listing_subset)
								tab =document.add_table(rows=row_body_cnt+row_header_cnt,cols=1+col_header_cnt)  

								#merge cell in header
								if header.startswith('file'):
									default=1
								else:
									default=0	  
								cell_merge_list,header_firstline_cell_list,header_lastline_cell_list=header_cellstyle_set(header_list,tab,default)   
								start=header_lastline_cell_list[-1]+(row_body_cnt-1)*(col_header_cnt+1)+1
								end=header_lastline_cell_list[-1]+row_body_cnt*(col_header_cnt+1)+1
								body_lastline_cell_list=[s for s in range(start,end)]

								#set the border
								border_set(tab,header_firstline_cell_list,cell_merge_list,header_lastline_cell_list,body_lastline_cell_list)

								#set value for header, and center the cell
								set_value_header(tab,header_list)
									#set value the first column for heaer
								if default==0:
									p=tab.cell(row_header_cnt-1,0).paragraphs[0]
									p.paragraph_format.space_before=Pt(2)	
									p.paragraph_format.space_after=Pt(2)	   
									p.paragraph_format.line_spacing=Pt(8)	   
									first_col_hd_tolist=ast.literal_eval(first_col_hd)
									cell_set(p,first_col_hd_tolist)
								else:
									first_col_hd_tolist=ast.literal_eval(first_col_hd_default[0])
									len_first_col_hd=len(first_col_hd_tolist)
									for i,first_col in enumerate(first_col_hd_tolist):
										p=tab.cell(row_header_cnt-len_first_col_hd+i,0).paragraphs[0]
										p.paragraph_format.space_before=Pt(2)	
										p.paragraph_format.space_after=Pt(2)	   
										p.paragraph_format.line_spacing=Pt(8)	  
										if first_col[0]!=[]:
											cell_set(p,first_col[0])
								#set value for body
								row_num=row_header_cnt  
								for si,row in enumerate(list(df_listing_subset['row'])):
									row_list=ast.literal_eval(row)
									len_row=len(row_list)
									for j,col in enumerate(body_lastline_cell_list):
										p=tab.cell(row_num+si,j).paragraphs[0]
										p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
										p.paragraph_format.space_before=Pt(2)	 
										p.paragraph_format.space_after=Pt(2)	  
										p.paragraph_format.line_spacing=Pt(8)	  
										if j==0:
											cell_liststyle=row_list[j][1:][1]
										elif j<len_row:
											cell_liststyle=row_list[j][1]
										else:
											cell_liststyle=row_list[len_row-1][1]					
										cell_set(p,cell_liststyle) 
									
							#set value for footnote
							footnote_style=df_footnote_subset.iloc[0,3]
							footnote_style=ast.literal_eval(footnote_style)
							p=document.add_paragraph()
							bullet_flag=0
							if len(footnote_style)>0:
								for m in range(len(footnote_style)):
									if len(footnote_style[m])>0:
										if len(footnote_style[m][1])>0 and footnote_style[m][0].startswith('BeiGene')==0 and footnote_style[m][0].startswith('Protocol BGB')==0: #check whether footnote is null
											if bullet_flag==1 and footnote_style[m][0].strip()!='':
												footnote_style[m][1][0][0]='  ●  '+footnote_style[m][1][0][0]
											cell_set(p,footnote_style[m][1]) 
											if footnote_style[m][0].strip().upper().startswith('PROGRAMMING NOTE'):
												bullet_flag=1 
									p.add_run('\n')   
							if tbl_flag==1: 
								document.add_page_break() 

					except Exception as e:
						print(traceback.format_exc())
						self.m_statusBar2.SetStatusText("Finding Error While Creating Listing From Excel: "+table_L.iloc[i,9])
					
					else:
				
						try:
							#set value for figure
							for i in range(len(table_F)):   

								self.m_statusBar2.SetStatusText("Finding Error While Creating Figure From Excel: "+table_F.iloc[i,9])  

								#check whether picture exists
								if table_F.iloc[i,2]!=table_F.iloc[i,2]:
									tbl_flag=0
								else:
									tbl_flag=1

								table_num=table_F.iloc[i,0]
								footnote=table_F.iloc[i,1]
								figure_code=table_F.iloc[i,2]
								df_footnote_subset=df_footnote[df_footnote.table_num==table_num]	  

								#read title, contain：table number，title，population
								title_set(table_F)

								if tbl_flag==1:
									# p=document.add_paragraph()
									document.add_picture(fig_read_loc+'pic'+figure_code+'.jpg')
									paragraph_containing_picture = document.paragraphs[-1]
									paragraph_containing_picture.alignment = WD_ALIGN_PARAGRAPH.CENTER

								#set value for footnote
								footnote_style=df_footnote_subset.iloc[0,3]
								footnote_style=ast.literal_eval(footnote_style)
								p=document.add_paragraph()
								bullet_flag=0
								if len(footnote_style)>0:
									for m in range(len(footnote_style)):
										if len(footnote_style[m])>0:
											if len(footnote_style[m][1])>0 and footnote_style[m][0].startswith('BeiGene')==0 and footnote_style[m][0].startswith('Protocol BGB')==0: #check whether footnote is null
												if bullet_flag==1 and footnote_style[m][0].strip()!='':
													footnote_style[m][1][0][0]='  ●  '+footnote_style[m][1][0][0]
												cell_set(p,footnote_style[m][1]) 
												if footnote_style[m][0].strip().upper().startswith('PROGRAMMING NOTE'):
													bullet_flag=1
										p.add_run('\n')   
								if tbl_flag==1: 
									document.add_page_break() 

						except Exception as e:
							print(traceback.format_exc())
							self.m_statusBar2.SetStatusText("Finding Error While Creating Figure From Excel: "+table_F.iloc[i,9])  

						else:

							document.save(output_loc)

							# 设置页眉页脚内容--除了页码
							# document = Document(output_loc)
							# styles = document.styles
							# style = styles.add_style("Header", WD_STYLE_TYPE.PARAGRAPH)
							# style.base_style = styles["Normal"]
							# tab_stops = style.paragraph_format.tab_stops
							# tab_stops.add_tab_stop(Inches(8.95), WD_TAB_ALIGNMENT.RIGHT)
							# sec_header = section.header
							# sec_footer = section.footer
							# sec_p=sec_header.paragraphs[0]
							# sec_p.add_run().add_picture(pic_loc,width=Inches(1.09))
							# sec_p.add_run("\t"+str(Protocol)+"\tTables, Figures, Listings (TFL) Shells")
							# sec_p.paragraph_format.space_before=Pt(15)	 #设置段前间距
							# sec_p.paragraph_format.space_after=Pt(1)	   #设置段后间距
							# sec_p.paragraph_format.line_spacing=Pt(10)	   #设置行间距
							# sec_p.style=document.styles['Header']
							# sec_f=sec_footer.paragraphs[0]
							# sec_f.add_run("Version: "+version+", "+draftdate)
							# document.save(output_loc)

							# #设置页码：
							# from win32com.client import Dispatch as MakeDoc
							# WordDoc = MakeDoc("Word.Application")
							# WordDoc = WordDoc.Documents.Add(output_loc)
							# WordDoc.Sections(1).Footers(1).PageNumbers.Add(2,True)
							# WordDoc.Sections(1).Footers(1).PageNumbers.NumberStyle = 57
							# WordDoc.SaveAs(output_loc)

			############################## End ######################################

							self.m_statusBar2.SetStatusText("The Process of Creating Mockup is Sucessfully End")
			else:
				self.m_statusBar2.SetStatusText("Please check the types of files extension")				


app = wx.App(False)
frame = MyFrame1(None)
frame.Show()
app.MainLoop()

