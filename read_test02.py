def read_headline(table):
    headline_flag=[10]
    for n, row in enumerate(table.rows):   # 读每行
        if n < 10:
            for j,cell in enumerate(row.cells):  # 读一行中的所有单元格
                if table.cell(n,0).text.strip() != '':
                    if table.cell(n,1).text.find('xx (') >= 0 or table.cell(n,1).text.strip()=='xx':
                        headline_flag.append(n-1)   
                if table.cell(n,0).text.strip() == '' or table.cell(n,0).text.strip() != '':
                    cellmiss=0
                    for i in range(1,len(row.cells)):
                        if table.cell(n,i).text.strip()=='':
                            cellmiss=cellmiss+1
                    if cellmiss==len(row.cells)-1:
                        headline_flag.append(n-1) 
                if table.cell(n,0).text.strip().startswith('xxx'):
                    headline_flag.append(n-1) 
    headline_flag_all.append(min(headline_flag))  
    return(headline_flag_all)

import docx
from docx import Document
import re
#mockup位置
#test_d0 = 'C:/Users/mingfeng.zhou/Desktop/standard_mockup/BeiGene non-efficacy TFL standard_02Sep2019_final_test4.docx'
# test_d1 = 'C:/Users/mingfeng.zhou/Desktop/standard_mockup/BeiGene non-efficacy TFL standard_02Sep2019_final_test1.docx'
# test_d2 = 'C:/Users/mingfeng.zhou/Desktop/standard_mockup/BeiGene non-efficacy TFL standard_02Sep2019_final_test1.docx'
test_d0 = 'C:/Users/mingfeng.zhou/Desktop/standard_mockup/BeiGene non-efficacy TFL standard_02Sep2019_final_test4.docx'
files=[test_d0]
files=files

#下载的图片本地保存位置
output_pictures='C:/Users/mingfeng.zhou/Desktop/standard_mockup/mockup-pictures/'
# output_pictures=file_list[0]+'/mockup-pictures/'

def read_headline(table):
    headline_flag=[10]
    for n, row in enumerate(table.rows):   # 读每行
        if n < 10:
            # for j,cell in enumerate(row.cells):  # 读一行中的所有单元格
            if table.cell(n,0).text.strip() != '':
                if table.cell(n,1).text.find('xx (') >= 0 or table.cell(n,1).text.strip()=='xx':
                    headline_flag.append(n-1)   
            if table.cell(n,0).text.strip() == '' or table.cell(n,0).text.strip() != '':
                cellmiss=0
                # print('---len.row.cells',len(row.cells))
                for i in range(1,len(row.cells)):
                    # if n<2:
                        # print('----cell.text---',table.cell(n,i).text.strip())
                    if table.cell(n,i).text.strip()=='':
                        cellmiss=cellmiss+1
                    # print('----null num--cellmiss',cellmiss)
                if cellmiss==len(row.cells)-1:
                    headline_flag.append(n-1) 
                    # print('--null all----',n-1)
            if table.cell(n,0).text.strip().startswith('xxx'):
                headline_flag.append(n-1) 
    return(min(headline_flag))


paragraphlist=[]
# 读取正文内容：Paragraphs
# print("段落数:"+str(len(document.paragraphs)))#段落数
table_num=[]
title=[]
population=[]
footnote=[]
table_code=[]
table_all=[]
footnote_sub=''
footnote_flag=0 #给footnote一个默认不存在的标志
fig_code=[]
footnote_sub_list=[]
table_contents_all=[]
first_col_hd=[]
first_col_hd_default=[]
param_all=[]
headline_flag_all=[]
table_hd_all=[]

for fn,test_d in enumerate(files):
    document=Document(test_d)
    table_value=10000
    datasource_num=100000
    for i,paragraph in enumerate(document.paragraphs):
        text= paragraph.text
        code=paragraph._p.getnext()
        paragraphlist.append(text)
        xmlcode=paragraph._p.xml
        run=paragraph.runs
        #print(i)
        #print('---text---',text)
        # print(code.tag)
        #print('---code---',code)

        if text.startswith(('Table','List','Figure')) and text.rstrip().endswith(')'): 
            table_value=i  
            if 'file'+str(fn)+'_'+text not in table_num:
                table_num.append('file'+str(fn)+'_'+text)
                paragraphlist[table_value]='file'+str(fn)+'_'+text
            else:
                table_num.append('file'+str(fn)+'_'+text+'01')
                paragraphlist[table_value]='file'+str(fn)+'_'+text+'01'                                    
            # print(paragraphlist[table_value])
            tbl_code=''
            rid='' 


            if footnote_flag==0 and len(table_num)==2: #当第一个table的footnote不存在时,将footnote赋值为空
                footnote=[[table_num[0],'\n',[['\n',[['\n',[None,None,None,None]]]]]]]
            elif footnote_flag==0 and len(table_num)>2: #当第>1个table的footnote不存在时,将footnote赋值为空
                footnote.append([table_num[-2],'\n',[['\n',[['\n',[None,None,None,None]]]]]])

            footnote_flag=0  #给footnote一个默认不存在的标志

        if 'graphicData' in xmlcode:  #查询图片位置
            rid=str(fn)+"".join(re.findall(r'a:blip r:embed="(\w+)"',xmlcode))
            headline_flag_all.append([paragraphlist[table_value],0])

        if i > table_value and i < table_value + 5 and code.tag.endswith('tbl'):
            tbl_code=code
            if tbl_code!='':
                for m,tbl in enumerate(document.tables):
                    if tbl._tbl==tbl_code: # 获取对应位置的table
                        print('----tbl_code----',tbl_code)


                        # self.m_statusBar2.SetStatusText("Reading Mockup "+paragraphlist[table_value]+" ...")

                        headline_flag=read_headline(tbl)
                        headline_flag_all.append([paragraphlist[table_value],headline_flag]) #获取header最后一行行号
                        row_text=[]
                        row_id=[]
                        row_hd=[]
                        first_col_hd_subset=[]
                        first_col_hd_subset_default=[]
                        for n, row in enumerate(tbl.rows):   # 读每行
                            col_text=[]
                            col_hd=[]
                            first_hd_0=[]
                            for j,cell in enumerate(row.cells):  # 读一行中的所有单元格
                                cell_run_text=[]
                                for z,cell_runs in enumerate(cell.paragraphs):
                                    cell_run = cell_runs.runs
                                    for k,cell_p in enumerate(cell_run):
                                        p_font=[cell_p.text,[cell_p.bold,cell_p.italic,cell_p.underline,cell_p.font.superscript]]   
                                        cell_run_text.append(p_font)
                                    if z<len(cell.paragraphs)-1:
                                        cell_run_text.append(['\n',[None,None,None,None]])

                                if j>0 and 0<n<=headline_flag: #第>1列单元格下一行等于上一行的值时，将下一行设为空
                                    # print('-------before compare-----------',tbl.cell(n,j).text,tbl.cell(n-1,j).text)
                                    if tbl.cell(n,j).text==tbl.cell(n-1,j).text and tbl.cell(n,j).text.strip()!="":
                                        # print('------------',tbl.cell(n,j).text)
                                        cell.text='\n'
                                        cell_run_text=[ ['\n',[None,None,None,None]] ]

                                if j > 0 :
                                    col_hd.append([cell.text,cell_run_text])

                                if j==0 and n>0: #第一列单元格下一行等于上一行的值时，将下一行设为空
                                    if tbl.cell(n,0).text==tbl.cell(n-1,0).text and tbl.cell(n,0).text.strip()!="":
                                        cell.text='\n'
                                        cell_run_text=[ ['\n',[None,None,None,None]] ]

                                if j ==0:
                                    col_text.append([paragraphlist[table_value],cell.text,cell_run_text])
                                    first_hd_0.append(cell_run_text)
            
                                else:
                                    col_text.append([cell.text,cell_run_text])  

                                if n<=headline_flag and j==0:
                                    first_col_hd_subset.extend(cell_run_text)
                                    # if n<headline_flag:
                                    #	 first_col_hd_subset=first_col_hd_subset.extend([['\n',None,None,None,None]])									   
    
                            # if row.cells[0].text.strip()!='':
                            row_text.append(col_text)
                        
                            if n <= headline_flag:
                                colmiss=0
                                for ms in range(0,len(col_hd)):
                                    if col_hd[ms][0].strip()=='':
                                        colmiss=colmiss+1
                                if colmiss!=len(col_hd):
                                    row_hd.append(col_hd)
                                    first_col_hd_subset_default.append(first_hd_0)

                            if len(row_text[n])>0:
                                if len(row_text[n][0]) > 0 and n > headline_flag:
                                    if len(row_text[n][0][2])>0:
                                        if row_text[n][0][2][0][0].lstrip()==row_text[n][0][2][0][0] and paragraphlist[table_value].startswith('file'+str(fn)+'_Table'):
                                            param_all.append([paragraphlist[table_value],row_text[n][0][1],row_text[n][0][2]])
                        table_contents_all.append(row_text)   
                        # if paragraphlist[table_value].find('_Listing')==-1:
                        table_hd_all.append([paragraphlist[table_value],row_hd])
                        first_col_hd.append([paragraphlist[table_value],first_col_hd_subset])
                        first_col_hd_default.append([paragraphlist[table_value],first_col_hd_subset_default])
        run_text=[]
        for p in run:
            p_font=[p.text,[p.bold,p.italic,p.underline,p.font.superscript]]
            run_text.append(p_font)

        if i==table_value+1:
            title.append([paragraphlist[table_value],text,run_text])

        if i==table_value+2:
            population.append([paragraphlist[table_value],text,run_text])

        if text.strip().startswith('Data source:')  or text.strip().startswith('Refer to'):
            datasource_num=i
            protocol_flag=0
            footnote_flag=1 #由Data source判断footnote存在
            # print(datasource_num)
            footnote_sub=text
            footnote_sub_list=[[text,run_text]]
            # print('____foot_init_',footnote_sub_list)
            table_code.append(tbl_code)  
            fig_code.append([paragraphlist[table_value],rid])

        if i> datasource_num:
            footnote_sub_list.append([text,run_text])
            # print('append',footnote_sub_list)
            footnote_sub=footnote_sub+text

            if footnote_flag==1:
                if text.strip().startswith('Protocol'):
                    protocol_flag=1
                    #print(footnote_sub_list)
                    footnote_sub=footnote_sub.replace('\n'+paragraphlist[i-1],'').replace(paragraphlist[i-1],'').replace(text,'')
                    if footnote_sub_list[-2][0].strip().startswith('BeiGene'):
                        footnote_list=[paragraphlist[table_value],footnote_sub,footnote_sub_list[0:len(footnote_sub_list)-2]] 
                    else:
                        footnote_list=[paragraphlist[table_value],footnote_sub,footnote_sub_list[0:len(footnote_sub_list)-1]]
                    footnote.append(footnote_list)
                    #print(footnote)

                elif text.startswith(('Table','List','Figure')) and text.rstrip().endswith(')'):
                    if protocol_flag==0:		
                        #print(footnote_sub_list)
                        footnote_sub=footnote_sub.replace('\n'+paragraphlist[i-1],'').replace(paragraphlist[i-1],'').replace(text,'')
                        footnote_list=[table_num[-2],footnote_sub,footnote_sub_list[0:len(footnote_sub_list)-1]]
                        footnote.append(footnote_list)
                        #print(footnote)
    if footnote_flag==1 or fn==len(files)-1: #当读取到最后一个table/Listing/Figure时，自动补上
        footnote_list=[paragraphlist[table_value],footnote_sub,footnote_sub_list]
        footnote.append(footnote_list)

# 将document中的图片下载到本地
    for shape in document.inline_shapes:
        contentID = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        # print(contentID)
        contentType = document.part.related_parts[contentID].content_type
        # print (shape.height.cm,shape.width.cm,shape._inline.graphic.graphicData.pic.nvPicPr.cNvPr.name) 
        if not contentType.startswith('image'):
            continue
        imgName = document.part.related_parts[contentID].partname
        imgData = document.part.related_parts[contentID]._blob
        # print(imgName,imgData)
        with open(output_pictures+'pic'+str(fn)+contentID+'.jpg','wb') as fp:
            fp.write(imgData)
    if len(document.inline_shapes) > 0:
        fp.close()

import pandas as pd
import numpy as np 
import xlsxwriter
excel_loc='C:/Users/mingfeng.zhou/Desktop/standard_mockup/result1.xlsx'
# excel_loc=os.path.join(file_list[0],'MockupToExcel.xlsx')

df_title=pd.DataFrame(title,columns=['table_num','title','title_sep'])
df_population=pd.DataFrame(population,columns=['table_num','population','pop_sep'])
df_footnote=pd.DataFrame(footnote,columns=['table_num','footnote','footnote_style'])
df_first_col_hd=pd.DataFrame(first_col_hd,columns=['table_num','first_col_hd'])
df_first_col_hd_default=pd.DataFrame(first_col_hd_default,columns=['table_num','first_col_hd_default'])
df_fig_code=pd.DataFrame(fig_code,columns=['table_num','fig_code'])
df_table_hd_all=pd.DataFrame(table_hd_all,columns=['table_num','hd_contents'])

df_tab=pd.concat([df_title[['table_num','title']], df_population[['population']], 
            df_footnote[['footnote']],df_fig_code['fig_code']],
            axis=1)
# df_tab=pd.merge(df_tab,df_first_col_hd[['table_num','first_col_hd']],how='left',on=['table_num'])
# df_tab=pd.merge(df_tab,df_table_hd_all,how='left',on=['table_num'])
# df_tab['header']='default'
# df_tab['headerline_flag']=headline_flag_all
# df_tab=pd.merge(df_tab,df_first_col_hd_default[['table_num','first_col_hd_default']],how='left',on=['table_num'])

# variable_all=[]
# listing_all=[]
# non_param_all=[]
# for i,tables in enumerate(table_contents_all):
#     for j,row in enumerate(tables): 
#         print(i,j,row)
#         if len(row)>0 and j > headline_flag_all[i]:
#             if len(row[0]) > 0:
#                 if row[0][1].strip()!='' and row[0][0].strip().find('_Table')>=0:
#                     variable_all.append(row[0])
#                     # non_param_all.append([row[0][0],row])
#         if row[0][0].strip().find('_Listing')>=0 and j > headline_flag_all[i]:
#             listing_all.append([row[0][0],row])
#         if row[0][0].strip().find('_Table')>=0 and j > headline_flag_all[i]:
#             non_param_all.append([row[0][0],row])
# pass
# df_variable=pd.DataFrame(variable_all,columns=['table_num','variable','format_style'])
pass