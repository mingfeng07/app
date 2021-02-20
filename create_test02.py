from docx import Document
import pandas as pd
import numpy as np 
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT
import ast
# 读取header from new docx
def read_header_from_docx(path):
    document=Document(path)
    # doc_tb=Document()
    header_list=[]
    for paragraph in document.paragraphs:
        text= paragraph.text
        if 'Study' in text:
            header_list.append(text)
    header_contents_all=dict()
    single_cell_underline_all=dict()
    gap_col_all=dict()
    for m,tbl in enumerate(document.tables):
        row_text=[]
        single_cell_underline_list=[]
        gap_col_list=[]
        for n, row in enumerate(tbl.rows):   # 读每行
            col_text=[]
            for j,cell in enumerate(row.cells):  # 读一行中的所有单元格

                # read gap column
                if j < len(row.cells)-1 and j > 0 and n < len(tbl.rows)-1:
                    if tbl.rows[n].cells[j].text=='' and tbl.rows[n+1].cells[j].text=='':
                        if j==1 and tbl.rows[n].cells[j+1].text==tbl.rows[n].cells[j+2].text:
                            if j+1 not in gap_col_list:
                                gap_col_list.append(j+1)
                        elif j > 1:
                            if (tbl.rows[n].cells[j+1].text==tbl.rows[n].cells[j+2].text) or (tbl.rows[n].cells[j-2].text==tbl.rows[n].cells[j-1].text):
                                if j+1 not in gap_col_list:
                                    gap_col_list.append(j+1)
                
                cell_run_text=[]  
                for z,cell_runs in enumerate(cell.paragraphs):
                    cell_run = cell_runs.runs
                    for k,cell_p in enumerate(cell_run):
                        p_font=[cell_p.text,[cell_p.bold,cell_p.italic,cell_p.underline,cell_p.font.superscript]]   
                        cell_run_text.append(p_font)
                    if z==0 and len(cell.paragraphs)>1:
                        cell_run_text.append(['\n',[None,None,None,None]])
                col_text.append([cell.text,cell_run_text])							 
            row_text.append(col_text)
        header_contents_all[header_list[m]]=row_text

        #read gap column
        gap_col_all[header_list[m]]=gap_col_list
        print(gap_col_all)
        #read single cell underline from new docx
        tbl0=tbl._tbl 
        for i,cell in enumerate(tbl0.iter_tcs()):  
            tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
            tcBorders = OxmlElement('w:tcBorders')
            bottom = OxmlElement('w:bottom')
            underline=tcPr.first_child_found_in("w:tcBorders")
            if underline is not None:
                single_cell_underline_list.append(i)	
        single_cell_underline_all[header_list[m]]=single_cell_underline_list	 
    return (header_contents_all,single_cell_underline_all,gap_col_all)




#定义title,包含：table number，title，population
def title_set(tableset):
    run1=document.add_heading(tableset.iloc[i,9], level=1)
    run1.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1.paragraph_format.space_before=Pt(5)	 #设置段前间距
    run1.paragraph_format.space_after=Pt(5)	   #设置段后间距
    run1.paragraph_format.line_spacing=Pt(8)	   #设置行间距
    run2=document.add_heading(tableset.iloc[i,5], level=2)  
    run2.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER 
    run2.paragraph_format.space_before=Pt(5)	 #设置段前间距
    run2.paragraph_format.space_after=Pt(5)	   #设置段后间距
    run2.paragraph_format.line_spacing=Pt(8)	   #设置行间距
    run3=document.add_paragraph(tableset.iloc[i,6])
    run3.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER 
    run3.paragraph_format.space_before=Pt(5)	 #设置段前间距
    run3.paragraph_format.space_after=Pt(5)	   #设置段后间距
    run3.paragraph_format.line_spacing=Pt(8)	   #设置行间距

#控制单元格格式，并返回单元格合并列表与header最后一行列表
def header_cellstyle_set(header_list,tab,default):
    row_header_cnt=len(header_list)
    col_header_cnt=len(header_list[0]) 
    cell_merge_of_list=[]
    header_lastline_cell_list=[]
    header_firstline_cell_list=[]
    dict_value=0   
    for n in range(row_header_cnt):
        for k in range(col_header_cnt+1):
            if k==0 and n>0:
                if default==0:
                    row_up,row_down=tab.cell(n-1,k),tab.cell(n,k)
                    row_up.merge(row_down)
                dict_value=dict_value+1
            elif k>1:
                if header_list[n][k-2][0]==header_list[n][k-1][0] and n==0: 
                    col_lt,col_rt=tab.cell(n,k-1),tab.cell(n,k)
                    col_lt.merge(col_rt)	
                    cell_merge_of_list.append(dict_value)
                elif header_list[n][k-2][0]==header_list[n][k-1][0] and n>0:
                    if header_list[n-1][k-2][0]==header_list[n-1][k-1][0] and header_list[n-1][k-1][0].strip().upper()!='(N = XX)':			  
                        col_lt,col_rt=tab.cell(n,k-1),tab.cell(n,k)
                        col_lt.merge(col_rt)	
                        cell_merge_of_list.append(dict_value)						
                    elif n>1 and header_list[n-1][k-2][0]==header_list[n-1][k-1][0] and header_list[n-1][k-1][0].strip().upper()=='(N = XX)':
                        if header_list[n-2][k-2][0]==header_list[n-2][k-1][0]:  
                            col_lt,col_rt=tab.cell(n,k-1),tab.cell(n,k)
                            col_lt.merge(col_rt)	
                            cell_merge_of_list.append(dict_value)							
                        else:
                            dict_value=dict_value+1
                    else:
                        dict_value=dict_value+1
                else:
                    dict_value=dict_value+1
            elif k==1:
                dict_value=dict_value+1
            if n==row_header_cnt-1:
                header_lastline_cell_list.append(dict_value)
            if n==0:
                header_firstline_cell_list.append(dict_value)
    return(cell_merge_of_list,header_firstline_cell_list,header_lastline_cell_list)

#设置单元格边界线
def border_set(tab,header_firstline,header_mid,header_lastline,body_lastline):
    tbl=tab._tbl 
    for i,cell in enumerate(tbl.iter_tcs()):
        if i in header_firstline:
            tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
            tcBorders = OxmlElement('w:tcBorders')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'thick')
            tcBorders.append(top)
            tcPr.append(tcBorders)				
        if i in header_firstline and i in header_mid:
            tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
            tcBorders = OxmlElement('w:tcBorders')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'thick')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'thick')			
            tcBorders.append(top)
            tcBorders.append(bottom)
            tcPr.append(tcBorders)
        elif i in header_mid or i in header_lastline or i in body_lastline:
            tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
            tcBorders = OxmlElement('w:tcBorders')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'thick')
            tcBorders.append(bottom)
            tcPr.append(tcBorders)			  

#set value and style for cell:
def cell_set(p,cell_list):
    for m in range(len(cell_list)):
        run=p.add_run(cell_list[m][0])
        run.bold=cell_list[m][1][0]
        run.italic=cell_list[m][1][1]
        run.underline=cell_list[m][1][2]
        run.font.superscript=cell_list[m][1][3]	 


#set value for header cell except for first column header:
def set_value_header(tab,header_list):
    row=len(header_list)
    col=len(header_list[0])
    for i in range(row):
        for j in range(col):		  
            cell_list=header_list[i][j][1]
            if j>0:
                if header_list[i][j][0]!=header_list[i][j-1][0]:	
                # if header_list[i][j]!=header_list[i][j-1]:	  
                    p=tab.cell(i,j+1).paragraphs[0]
                    cell_set(p,cell_list)
                    # p.add_run(cell_list)
                elif header_list[i][j][0]==header_list[i][j-1][0] and (header_list[i][j][0].strip().upper()=='(N = XX)' or header_list[i][j][0].strip().upper()=='n/N (%)'):
                    if i>0:
                        if header_list[i-1][j][0]!=header_list[i-1][j-1][0]:
                            p=tab.cell(i,j+1).paragraphs[0]
                            cell_set(p,cell_list)
                            # p.add_run(cell_list)
                elif header_list[i][j][0]==header_list[i][j-1][0] and header_list[i][j][0].strip()=='n (%)':
                    if i>0:
                        if header_list[i-1][j][0]!=header_list[i-1][j-1][0]:
                            p=tab.cell(i,j+1).paragraphs[0]
                            cell_set(p,cell_list)
                            # p.add_run(cell_list)							 
                    if i>1:
                        if header_list[i-1][j][0]==header_list[i-1][j-1][0] and header_list[i-1][j][0].strip().upper()=='(N = XX)':
                            if header_list[i-2][j][0]!=header_list[i-2][j-1][0]:
                                p=tab.cell(i,j+1).paragraphs[0]
                                cell_set(p,cell_list)
                                # p.add_run(cell_list)					
            else:
                p=tab.cell(i,j+1).paragraphs[0]
                cell_set(p,cell_list)
                # p.add_run(cell_list)
            p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before=Pt(2)	 #设置段前间距
            p.paragraph_format.space_after=Pt(2)	   #设置段后间距
            p.paragraph_format.line_spacing=Pt(8)	   #设置行间距

#定义通用变量
header_loc = 'C:/Users/mingfeng.zhou/Desktop/standard_mockup/header_for_tables.docx'
# header_loc=word_str
mp_excel_loc='C:/Users/mingfeng.zhou/Desktop/standard_mockup/MockupToExcel.xlsx'
# mp_excel_loc=excel_str
# pic_loc= 'C:/Users/mingfeng.zhou/Desktop/standard_mockup/BeiGene_pic.png'
fig_read_loc= 'C:/Users/mingfeng.zhou/Desktop/standard_mockup/mockup-pictures/'
# fig_read_loc=file_list[0]+'/mockup-pictures/'
output_loc='C:/Users/mingfeng.zhou/Desktop/standard_mockup/demo_output.docx'
# output_loc=os.path.join(file_list[0],'Mockup-Output.docx')
# Protocol="BGB-3111-210"
# version="1.6"
# draftdate="12 Aug 2019"

doc_header=Document(header_loc)
document = Document()

sections=document.sections
section = document.sections[-1]
#文档横向设置：
section.orientation = WD_ORIENT.LANDSCAPE
section.page_height,section.page_width=Inches(8.5),Inches(11)
#section.left_margin, section.right_margin，section.top_margin, section.bottom_margin，section.header_distance, section.footer_distance，section.page_width, section.page_height

#正文字体和尺寸控制：
document.styles['Normal'].font.name = u'Courier New'   
document.styles['Normal'].font.size=Pt(8)

#控制标题格式：
document.styles['Heading 1'].font.size = Pt(8)
document.styles['Heading 1'].font.name = u'Courier New'
document.styles['Heading 1'].font.color.rgb = None
document.styles['Heading 1'].font.italic = False
document.styles['Heading 1'].font.bold = False
document.styles['Heading 2'].font.size = Pt(8)
document.styles['Heading 2'].font.name = u'Courier New'
document.styles['Heading 2'].font.color.rgb = None
document.styles['Heading 2'].font.italic = False
document.styles['Heading 2'].font.bold = False

table_ori=pd.read_excel(mp_excel_loc,sheet_name='table')
table_mp=pd.read_excel(mp_excel_loc,sheet_name='mockup')
table_mp=table_mp[table_mp.order>=0].sort_values('order')
table=pd.merge(table_ori[['table_num','footnote','fig_code','first_col_hd','hd_contents']],
                table_mp[['table_num','title','population','header','order','new_table_num']],
                how='right',on=['table_num'],sort=False)

#合并源表header与新docx的header
header_list_all,single_cell_underline_all,gap_col_all=read_header_from_docx(header_loc)
for i in range(len(table)):
    if table.iloc[i,7]=='default' and table.iloc[i,0].find('_Fig')==-1:
        table.iloc[i,7]=table.iloc[i,0]
        header_list_all[table.iloc[i,0]]=ast.literal_eval(table.iloc[i,4])

df_param=pd.read_excel(mp_excel_loc,sheet_name='param_style')
df_non_param=pd.read_excel(mp_excel_loc,sheet_name='non_param_style')
df_format=pd.read_excel(mp_excel_loc,sheet_name='format_style')
df_footnote=pd.read_excel(mp_excel_loc,sheet_name='footnote_style')
table_T=table[table.table_num.str.find('_Table')>-1]
table_F=table[table.table_num.str.find('_Fig')>-1]
table_L=table[table.table_num.str.find('_List')>-1]
df_listing=pd.read_excel(mp_excel_loc,sheet_name='listing_style')




#给Table赋值
for i in range(len(table_T)):

    # self.m_statusBar2.SetStatusText("Creating Table From Excel: "+table_T.iloc[i,9])

    table_num=table_T.iloc[i,0]
    header=table_T.iloc[i,7]
    header_list=header_list_all[header]
    first_col_hd=table_T.iloc[i,3]
    first_col_hd_default=list(table_ori[table_ori['table_num']==table_num]['first_col_hd_default'])
    footnote=table_T.iloc[i,1]
    df_param_subset=df_param[df_param.table_num==table_num]
    df_format_subset=df_format[df_format.table_num==table_num]
    df_footnote_subset=df_footnote[df_footnote.table_num==table_num]
    df_non_param_subset=df_non_param[df_non_param.table_num==table_num]

    #读入title,包含：table number，title，population
    title_set(table_T)

    row_header_cnt=len(header_list)
    col_header_cnt=len(header_list[0])
    row_body_cnt=len(df_non_param_subset)
    tab =document.add_table(rows=row_body_cnt+row_header_cnt,cols=1+col_header_cnt)  
        
    #控制header单元格合并   
    if header.startswith('file'):
        default=1
        single_cell_underline_list=[]
        gap_col_list=[]
    else:
        default=0
        single_cell_underline_list=single_cell_underline_all[header] 
        gap_col_list=gap_col_all[header]

    cell_merge_list,header_firstline_cell_list,header_lastline_cell_list=header_cellstyle_set(header_list,tab,default)   

    print('cellmerge=======',cell_merge_list)
    print('cellsingle=======',single_cell_underline_list)	
    print('cell_header_lastline====',header_lastline_cell_list)
    start=header_lastline_cell_list[-1]+(row_body_cnt-1)*(col_header_cnt+1)+1
    end=header_lastline_cell_list[-1]+row_body_cnt*(col_header_cnt+1)+1
    body_lastline_cell_list=[s for s in range(start,end)]

    #控制边界线
    border_set(tab,header_firstline_cell_list,cell_merge_list,header_lastline_cell_list,body_lastline_cell_list)

    #对header赋值，并控制单元格居中
    set_value_header(tab,header_list)
        #header第一列赋值
    if default==0:
        p=tab.cell(row_header_cnt-1,0).paragraphs[0]
        p.paragraph_format.space_before=Pt(2)	 #设置段前间距
        p.paragraph_format.space_after=Pt(2)	   #设置段后间距
        p.paragraph_format.line_spacing=Pt(8)	   #设置行间距
        first_col_hd_tolist=ast.literal_eval(first_col_hd)
        cell_set(p,first_col_hd_tolist)
    else:
        first_col_hd_tolist=ast.literal_eval(first_col_hd_default[0])
        len_first_col_hd=len(first_col_hd_tolist)
        for i,first_col in enumerate(first_col_hd_tolist):
            p=tab.cell(row_header_cnt-len_first_col_hd+i,0).paragraphs[0]
            p.paragraph_format.space_before=Pt(2)	 #设置段前间距
            p.paragraph_format.space_after=Pt(2)	   #设置段后间距
            p.paragraph_format.line_spacing=Pt(8)	   #设置行间距
            if first_col[0]!=[]:
                cell_set(p,first_col[0])
    #对body赋值
    row_num=row_header_cnt	
    for si,row in enumerate(list(df_non_param_subset['row'])):
        row_list=ast.literal_eval(row)
        len_row=len(row_list)
        for j,col in enumerate(body_lastline_cell_list):
            p=tab.cell(row_num+si,j).paragraphs[0]
            if j==0:
                p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER				
            p.paragraph_format.space_before=Pt(2)	 #设置段前间距
            p.paragraph_format.space_after=Pt(2)	   #设置段后间距
            p.paragraph_format.line_spacing=Pt(8)	   #设置行间距
            if j==0:
                cell_liststyle=row_list[j][1:][1]
            elif j<len_row:
                cell_liststyle=row_list[j][1]
            else:
                cell_liststyle=row_list[len_row-1][1]					
            cell_set(p,cell_liststyle)  
            # gap_column set null
            if j in gap_col_list:
                tab.cell(row_num+si,j).text=''

    #列控制宽度
    tab.autofit = False
    col = tab.columns[0] 

    #gap width set as 0.2inches
    gap_width=0.2
    gap_cnt=len(gap_col_list)

    col_cnt=col_header_cnt+1
    if col_cnt <=4:
        width_col0=4
    else:
        width_col0=4-(col_cnt-4)*0.5 
    width_col_oth=(8.5-width_col0-gap_width*gap_cnt)/(col_cnt-1-gap_cnt)
    for ilen in range(col_cnt):
        col = tab.columns[ilen] 
        if ilen==0:
            for cell in col.cells:
                cell.width = Inches(width_col0)
        elif ilen in gap_col_list:
            for cell in col.cells:
                cell.width = Inches(gap_width)
        else:
            for cell in col.cells:
                cell.width = Inches(width_col_oth)		   

    #对footnote赋值
    footnote_style=df_footnote_subset.iloc[0,3]
    footnote_style=ast.literal_eval(footnote_style)
    p=document.add_paragraph()
    if len(footnote_style)>0:
        for m in range(len(footnote_style)):
            if len(footnote_style[m])>0:
                if len(footnote_style[m][1])>0 and footnote_style[m][0].startswith('BeiGene')==0 and footnote_style[m][0].startswith('Protocol BGB')==0: #检查footnote是否为空
                    cell_set(p,footnote_style[m][1])  
            p.add_run('\n')   
    document.add_page_break() 
pass

