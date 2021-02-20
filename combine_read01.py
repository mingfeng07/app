#################################### Start Read Mockup ##################################								
				#  解决header的格式问题：  ##
				# 修改header，如果为default，则不合并第一列header单元格 #

				#识别header结束行并获取行号list
def read_headline(table):
    headline_flag=[10]
    for n, row in enumerate(table.rows):   # 读每行
        if n < 10:
            for j,cell in enumerate(row.cells):  # 读一行中的所有单元格
                if table.cell(n,0).text.strip() != '':
                    if table.cell(n,1).text.find('xx (') >= 0 or table.cell(n,1).text.strip()=='xx':
                        headline_flag.append(n-1)   
                if table.cell(n,0).text.strip() == '' or table.cell(n,0).text.strip() != '':
                    cellmiss=0
                    for i in range(1,len(row.cells)):
                        if table.cell(n,i).text.strip()=='':
                            cellmiss=cellmiss+1
                    if cellmiss==len(row.cells)-1:
                        headline_flag.append(n-1) 
                if table.cell(n,0).text.strip().startswith('xxx'):
                    headline_flag.append(n-1) 
    headline_flag_all.append(min(headline_flag))  
    return(headline_flag_all)

import docx
from docx import Document
import re
#mockup位置

test_d0 = 'C:/Users/mingfeng.zhou/Desktop/combine rtf/t-14-2-1-2-1-ef-dor-com-irc.docx'

files=[test_d0]
# files=files

#下载的图片本地保存位置
output_pictures='C:/Users/mingfeng.zhou/Desktop/standard_mockup/mockup-pictures/'
# output_pictures=file_list[0]+'/mockup-pictures/'

paragraphlist=[]
# 读取正文内容：Paragraphs
# print("段落数:"+str(len(document.paragraphs)))#段落数
table_num=[]
title=[]
population=[]
footnote=[]
table_code=[]
table_all=[]
footnote_sub=''
footnote_flag=0
fig_code=[]
footnote_sub_list=[]
table_contents_all=[]
first_col_hd=[]
first_col_hd_default=[]
param_all=[]
headline_flag_all=[]
table_hd_all=[]
for fn,test_d in enumerate(files):
    document=Document(test_d)
    table_value=10000
    datasource_num=100000
    for i,paragraph in enumerate(document.paragraphs):
        text= paragraph.text
        code=paragraph._p.getnext()
        paragraphlist.append(text)
        xmlcode=paragraph._p.xml
        run=paragraph.runs

        # print(i)
        print(code.tag)
        print(code)
        print(xmlcode)
pass

        # if paragraph.style.name=='List Paragraph':
        #     print('text.style---------',paragraph.style.name)
        if text.startswith(('Table','List','Figure')) and text.rstrip().endswith(')'): 
            table_num.append('file'+str(fn)+'_'+text)
            table_value=i  
            paragraphlist[table_value]='file'+str(fn)+'_'+text
            # print(paragraphlist[table_value])
            tbl_code=''
            rid='' 

        if 'Programming' in text:
            print('-----text---',text.strip().upper())

            # print(footnote,footnote_flag,len(table_num))

            if footnote_flag==0 and len(table_num)==2: #当第一个table的footnote不存在时,将footnote赋值为空
                footnote=[[table_num[0],'\n',[['\n',[['\n',[None,None,None,None]]]]]]]
            elif footnote_flag==0 and len(table_num)>2: #当第>1个table的footnote不存在时,将footnote赋值为空
                footnote[-1]=[table_num[-2],'\n',[['\n',[['\n',[None,None,None,None]]]]]]

        if i > table_value and i < table_value + 5 and code.tag.endswith('tbl'):
            tbl_code=code
            if tbl_code!='':
                for m,tbl in enumerate(document.tables):
                    if tbl._tbl==tbl_code: # 获取对应位置的table

                        footnote_flag=0  #给footnote一个默认不存在的标志

                        # self.m_statusBar2.SetStatusText("Reading Mockup "+paragraphlist[table_value]+" ...")

                        # headline_flag_all=read_headline(tbl) #获取header最后一行行号

                        # print(paragraphlist[table_value],headline_flag_all)
                        row_text=[]
                        row_id=[]
                        row_hd=[]
                        first_col_hd_subset=[]
                        first_col_hd_subset_default=[]
                        for n, row in enumerate(tbl.rows):   # 读每行
                            col_text=[]
                            col_hd=[]
                            first_hd_0=[]
                            for j,cell in enumerate(row.cells):  # 读一行中的所有单元格
                                print(cell.text)
pass

                                cell_run_text=[]
                                for z,cell_runs in enumerate(cell.paragraphs):
                                    cell_run = cell_runs.runs
                                    for k,cell_p in enumerate(cell_run):
                                        p_font=[cell_p.text,[cell_p.bold,cell_p.italic,cell_p.underline,cell_p.font.superscript]]   
                                        cell_run_text.append(p_font)
                                    if z<len(cell.paragraphs)-1:
                                        cell_run_text.append(['\n',[None,None,None,None]])

                                if j > 0 :
                                    col_hd.append([cell.text,cell_run_text])

                                if j ==0:
                                    col_text.append([paragraphlist[table_value],cell.text,cell_run_text])
                                    first_hd_0.append(cell_run_text)
            
                                else:
                                    col_text.append([cell.text,cell_run_text])  

                                if n<=headline_flag_all[-1] and j==0:
                                    first_col_hd_subset.extend(cell_run_text)
                                    # if n<headline_flag_all[-1]:
                                    #	 first_col_hd_subset=first_col_hd_subset.extend([['\n',None,None,None,None]])									   
    
                            # if row.cells[0].text.strip()!='':
                            row_text.append(col_text)
                        
                            if n <= headline_flag_all[-1]:
                                colmiss=0
                                for ms in range(0,len(col_hd)):
                                    if col_hd[ms][0].strip()=='':
                                        colmiss=colmiss+1
                                if colmiss!=len(col_hd):
                                    row_hd.append(col_hd)
                                    first_col_hd_subset_default.append(first_hd_0)

                            if len(row_text[n])>0:
                                if len(row_text[n][0]) > 0 and n > headline_flag_all[-1]:
                                    if len(row_text[n][0][2])>0:
                                        if row_text[n][0][2][0][0].lstrip()==row_text[n][0][2][0][0] and paragraphlist[table_value].startswith('file'+str(fn)+'_Table'):
                                            param_all.append([paragraphlist[table_value],row_text[n][0][1],row_text[n][0][2]])
                        table_contents_all.append(row_text)   
                        # if paragraphlist[table_value].find('_Listing')==-1:
                        table_hd_all.append([paragraphlist[table_value],row_hd])
                        first_col_hd.append([paragraphlist[table_value],first_col_hd_subset])
                        first_col_hd_default.append([paragraphlist[table_value],first_col_hd_subset_default])
        run_text=[]
        for p in run:
            p_font=[p.text,[p.bold,p.italic,p.underline,p.font.superscript]]
            run_text.append(p_font)

        if i==table_value+1:
            title.append([paragraphlist[table_value],text,run_text])

        if i==table_value+2:
            population.append([paragraphlist[table_value],text,run_text])

        if text.strip().startswith('Data source:'):
            datasource_num=i
            protocol_flag=0

            footnote_flag=1 #由Data source判断footnote存在

            # print(datasource_num)
            footnote_sub=text
            footnote_sub_list=[[text,run_text]]
            # print('____foot_init_',footnote_sub_list)
            table_code.append(tbl_code)  
            fig_code.append([paragraphlist[table_value],rid])

        if i> datasource_num:
            footnote_sub_list.append([text,run_text])
            # print('append',footnote_sub_list)
            footnote_sub=footnote_sub+text
            
            if text.strip().startswith('Protocol'):
                protocol_flag=1
                #print(footnote_sub_list)
                footnote_sub=footnote_sub.replace('\n'+paragraphlist[i-1],'').replace(paragraphlist[i-1],'').replace(text,'')
                # print('Protocol------------',footnote_sub)
                if footnote_sub_list[-2][0].strip().startswith('BeiGene'):
                    footnote_list=[paragraphlist[table_value],footnote_sub,footnote_sub_list[0:len(footnote_sub_list)-2]]     
                else:               
                    footnote_list=[paragraphlist[table_value],footnote_sub,footnote_sub_list[0:len(footnote_sub_list)-1]]
                footnote.append(footnote_list)
                #print(footnote)

            elif text.strip().startswith('Table ') or text.strip().startswith('Listing ') or text.strip().startswith('Figure '):
                if protocol_flag==0:
                    #print(footnote_sub_list)
                    footnote_sub=footnote_sub.replace('\n'+paragraphlist[i-1],'').replace(paragraphlist[i-1],'').replace(text,'')
                    footnote_list=[table_num[-2],footnote_sub,footnote_sub_list[0:len(footnote_sub_list)-1]]
                    footnote.append(footnote_list)
                    # print(footnote_sub_list[-1])
    footnote_list=[paragraphlist[table_value],footnote_sub,footnote_sub_list]
    footnote.append(footnote_list)

    #填补headline_flag_all的figure部分的值，默认为0
    len_headline_flag=len(headline_flag_all)
    for i in range(len(table_num)):
        if i >=len_headline_flag:
            headline_flag_all.append(0)

import pandas as pd
import numpy as np 
import xlsxwriter

excel_loc='C:/Users/mingfeng.zhou/Desktop/standard_mockup/result1.xlsx'
# excel_loc=os.path.join(file_list[0],'MockupToExcel.xlsx')

df_title=pd.DataFrame(title,columns=['table_num','title','title_sep'])
df_population=pd.DataFrame(population,columns=['table_num','population','pop_sep'])
df_footnote=pd.DataFrame(footnote,columns=['table_num','footnote','footnote_style'])
df_first_col_hd=pd.DataFrame(first_col_hd,columns=['table_num','first_col_hd'])
df_first_col_hd_default=pd.DataFrame(first_col_hd_default,columns=['table_num','first_col_hd_default'])
df_fig_code=pd.DataFrame(fig_code,columns=['table_num','fig_code'])
df_table_hd_all=pd.DataFrame(table_hd_all,columns=['table_num','hd_contents'])

df_tab=pd.concat([df_title[['table_num','title']], df_population[['population']], 
            df_footnote[['footnote']],df_fig_code['fig_code']],
            axis=1)
df_tab=pd.merge(df_tab,df_first_col_hd[['table_num','first_col_hd']],how='left',on=['table_num'])
df_tab=pd.merge(df_tab,df_table_hd_all,how='left',on=['table_num'])
df_tab['header']='default'
df_tab['headerline_flag']=headline_flag_all
df_tab=pd.merge(df_tab,df_first_col_hd_default[['table_num','first_col_hd_default']],how='left',on=['table_num'])


variable_all=[]
listing_all=[]
non_param_all=[]
for i,tables in enumerate(table_contents_all):
    for j,row in enumerate(tables): 

        if len(row)>0 and j > headline_flag_all[i]:
            if len(row[0]) > 0:
                if row[0][1].strip()!='' and row[0][0].strip().find('_Table')>=0:
                    variable_all.append(row[0])
                    # non_param_all.append([row[0][0],row])
        if row[0][0].strip().find('_Listing')>=0 and j > headline_flag_all[i]:
            listing_all.append([row[0][0],row])
        if row[0][0].strip().find('_Table')>=0 and j > headline_flag_all[i]:
            non_param_all.append([row[0][0],row])

df_variable=pd.DataFrame(variable_all,columns=['table_num','variable','format_style'])








