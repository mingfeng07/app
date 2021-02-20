import docx
from docx import Document
import re
#mockup位置

# test_d0 = 'C:/Users/mingfeng.zhou/Desktop/combine rtf/t-14-2-1-1-2-1-dor-inv-wm-eff.docx'
# test_d0 = 'C:/Users/mingfeng.zhou/Desktop/combine rtf/t-14-1-2-2-dh-saffl.docx'
# test_d0 = 'C:/Users/mingfeng.zhou/Desktop/combine rtf/t-14-2-1-2-1-ef-dor-com-irc.docx'
test_d0 = 'C:/Users/mingfeng.zhou/Desktop/combine rtf/test0001.docx'

files=[test_d0]
# files=files

#下载的图片本地保存位置
output_pictures='C:/Users/mingfeng.zhou/Desktop/standard_mockup/mockup-pictures/'
# output_pictures=file_list[0]+'/mockup-pictures/'

paragraphlist=[]
# 读取正文内容：Paragraphs
# print("段落数:"+str(len(document.paragraphs)))#段落数
table_num=[]
title=[]
population=[]
footnote=[]
table_code=[]
table_all=[]
footnote_sub=''
footnote_flag=0
fig_code=[]
footnote_sub_list=[]
table_contents_all=[]
first_col_hd=[]
first_col_hd_default=[]
param_all=[]
headline_flag_all=[]
table_hd_all=[]
for fn,test_d in enumerate(files):
    document=Document(test_d)
    for m,tbl in enumerate(document.tables):
        print('行数，列数',len(tbl.rows))
        for n, row in enumerate(tbl.rows):   # 读每行
            for j,cell in enumerate(row.cells):  # 读一行中的所有单元格
                print('行，列',n,j,'value=',cell.text)
pass